from pathlib import Path
from textwrap import dedent
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUT_MD = ROOT / "hybrid-engine-final-evidence-dossier.md"
OUT_DOCX = ROOT / "THE_Hybrid_Engine_Final_Evidence_Dossier.docx"

AUTHORED = dedent(r'''
# THE Hybrid Engine — Final Evidence Dossier and Implementation Lock

**Project status:** final close-out document  
**Research date:** 16 August 2026  
**Purpose:** settle the evidence questions, define the product-safe rules, and close the progression/deload research loop without pretending that engineering heuristics are experimentally proven constants.

## Executive verdict

This document closes the question. THE Hybrid Engine should be built as a transparent, local-first, deterministic adaptive training system. It should observe the athlete, estimate state with explicit uncertainty, preserve the purpose of the planned session where safe, and make bounded changes that are easy to audit. It should not be sold to the codebase or to the athlete as an artificial intelligence that has discovered the one true training formula.

The central research result is negative but useful: no convincing body of evidence was found that directly compares a progression rule of approximately 2.5% with 5% and 10% while holding exercise, population, trigger, volume, and context constant. The commonly repeated 2–10% recommendation is real as an ACSM position-stand recommendation, but the range is a practical prescription band with Category B support rather than an optimisation result. Plotkin et al. compared load progression with repetition progression; it did not test increment size. Hostler et al. showed that very small increments can work over eight weeks in upper-body exercises, but did not establish a universal optimum or a long-term adherence advantage. Buskard et al. found no clear strength or functional-performance winner among several progression methods in older adults. The evidence therefore supports a controller shape, not a magic constant.

The implementation decision is consequently firm but honestly labelled. The engine may use a default progression target of 2.5% of the last stable opening load, with equipment-aware rounding and a repetition/RPE fallback. That is an engineering default chosen conservatively inside the ACSM band. It is not a finding that 2.5% beats 5% or 10%. The engine may use a default 5% reduction from the last successful anchor after repeated, comparable deterioration. That is a product heuristic. It is not a validated physiological threshold, and the rule must not blindly compound a within-session failure correction.

The strongest product decision is more important than either percentage: do not let a single noisy observation control the whole programme. A progression decision should be based on comparable exposure, observed performance, exercise familiarity, technique, pain/symptom status, and the presence or absence of corroborating fatigue. A low HRV value, a bad night of sleep, or a subjective “flat” feeling may lower confidence, but none should independently authorise a heavier load or diagnose recovery. Pain must be represented separately from exertional fatigue. A training gap must enter a calibration state rather than trigger an invented percentage equation. Missing data must reduce certainty, not silently become normal data.

That is the final position: build the adaptive layer, but make its evidence status visible. The system’s intelligence should be its restraint, auditability, and ability to choose the next useful action when the evidence is incomplete.

## How to read this dossier

Every recommendation is tagged conceptually in one of four ways. **Direct evidence** means that a study tested the mechanism or population reasonably close to the proposed claim. **Adjacent evidence** means that the study informs the mechanism but does not validate the exact app rule. **Coaching precedent** means that a named system or experienced practice uses a similar idea; it can guide design but cannot be presented as causal proof. **Product-design convention** means that the number or state exists because the product needs a deterministic and conservative behaviour where the literature has not supplied one.

The dossier contains two layers. The first is the final interpretation and implementation lock written for the project. The second is the inherited evidence bundle and design record preserved as appendices. The appendices are included so that the decision can be audited later without reopening the same research from memory. They are not a claim that every sentence in an inherited note has equal evidential weight; the reconciliation chapters above take precedence where wording or certainty has changed.

## 1. The exact question the project needed answered

The practical problem was not whether progressive overload exists. It does. The problem was how a software engine should decide what to do after a successful or unsuccessful exposure when the user trains strength and conditioning in the same week, reports RPE/RIR, may have wearable recovery data, has finite equipment increments, and needs a plan that remains usable on bad real-world days.

That problem contains several questions which are often collapsed into one:

1. Does increasing external load generally support strength and hypertrophy?
2. Is increasing load better than increasing repetitions when the athlete is already working hard?
3. Does the size of the load increment matter, and is there a tested optimum?
4. Should the next increment be a fixed number of kilograms or a relative percentage?
5. How much evidence should be required before a progression change?
6. What should happen after one miss, repeated misses, or a whole-session decline?
7. When should a deload reduce load, volume, frequency, effort, or some combination?
8. How should conditioning fatigue, sleep, HRV, pain, illness, and a training gap alter the decision?
9. How can an app communicate uncertainty without turning it into a decorative confidence score?

The evidence is much better for the first two questions than for the third through ninth. The first two are training questions. The later questions are control-system questions wrapped around training. A study showing that a programme improves squat strength does not automatically tell us how many consecutive misses a mobile app should tolerate, what a plate-limited jump should be, or whether the app should reduce the next session by 5% or 7.5%. Those are separate claims with separate evidential burdens.

The project went wrong whenever it treated a familiar coaching number as though it were a measured natural constant. “Two consecutive misses,” “three sessions,” “5% deload,” and “10% lower-body progression” may each be reasonable starting conventions. None becomes scientifically established merely because several programmes use it. The final engine therefore records the provenance of every decision: evidence, precedent, heuristic, or safety policy.

## 2. Evidence method and standards of honesty

The search was organised around mechanisms rather than slogans. The core searches targeted the ACSM progression recommendation, the source cited behind it, direct comparisons of load and repetition progression, direct or near-direct evidence on microloading, autoregulated progression methods, deload studies and surveys, reliability and measurement error, HRV-guided resistance training, detraining and return to training, pain-monitoring models, and uncertainty communication. The search was then extended to named coaching and product systems because those systems are relevant precedents for how a practical feedback loop can be presented, even though their internal algorithms are not independent scientific validation.

The evidence was assessed with a deliberately conservative hierarchy. A position stand can be authoritative and still contain a recommendation whose precise dose has only limited experimental support. A randomised trial can be rigorous and still answer a narrower question than the product wants to ask. A meta-analysis can produce a pooled effect while inheriting heterogeneity in load, exercise, participants, and outcome measurement. A survey can describe what experienced athletes do while remaining unable to show that the practice caused their outcomes. A product’s help page can prove what the creator says the algorithm does, but not that the algorithm is optimal or generalisable.

The report therefore avoids three common errors. First, it does not convert absence of evidence into evidence of absence. “No direct trial found” means the reviewed search did not find one that supports the exact claim; it does not prove that no inaccessible dissertation or unpublished dataset exists. Second, it does not confuse a statistically significant result with a practically useful software rule. A tiny average difference may be statistically detectable and still not justify a different user-facing behaviour. Third, it does not hide population mismatch. Findings in untrained young men, older adults, trained powerlifters, or clinical tendinopathy patients are not automatically transferable to an unsupervised mixed-modality recreational user.

The engineering standard is therefore: every rule must state what it is allowed to claim, what it cannot claim, what inputs it requires, and what action it takes when its inputs are missing or contradictory. A rule that cannot explain itself is not ready for the engine, even if its number sounds familiar.

## 3. What ACSM supports—and what it does not

The 2009 ACSM position stand, “Progression Models in Resistance Training for Healthy Adults,” gives the familiar recommendation that when an athlete can perform one to two repetitions above the desired workload on two consecutive training sessions, the load may be increased by 2–10%. The paper also describes the lower end as more appropriate to smaller muscle-mass exercises and the higher end as more appropriate to larger, multi-joint exercises. This is directly relevant to the project because it validates the general direction of a relative, exercise-sensitive progression policy.

It does not validate a software constant. The sentence does not say that 2.5% is optimal for a bench press, that 5% is optimal for a squat, or that 10% is optimal for a deadlift. It does not provide a plate-rounding algorithm, a maximum acceptable rounded jump, or a fallback rule for when equipment resolution makes the target impossible. It does not establish that two sessions is superior to one or three. It is a practical prescription within an evidence category, not a dose-response experiment.

The source lineage matters. The 2009 position stand cites Feigenbaum and Pollock’s prescription paper for the progression recommendation. That source is a prescription and review contribution, not a randomised comparison of 2%, 5%, and 10% jumps. The recommendation is therefore downstream of accumulated exercise-prescription practice and expert synthesis. That does not make it useless. It does mean the code documentation must say “ACSM-supported operating range” rather than “trial-proven optimal percentage.”

The 2026 ACSM position stand strengthens the broader design philosophy. It synthesised 137 systematic reviews representing more than 30,000 participants and emphasised consistency, individualisation, goal-specific loading, and the fact that complicated techniques are not consistently necessary for the average healthy adult. It does not replace the missing direct increment-size experiment. Its value to this project is that it pushes the engine away from rigid one-size-fits-all rules and toward a system that is effective because the athlete can actually sustain it.

The product interpretation is precise:

- ACSM supports a progression band and a performance-based trigger.
- ACSM supports using smaller practical increases for smaller exercises and larger increases for larger exercises.
- ACSM does not supply the exact software constants.
- The 2026 update supports individualisation and adherence as first-order design requirements.
- The engine must preserve the uncertainty rather than laundering the recommendation into a false equation.

## 4. What Plotkin actually tested

Plotkin et al. compared two progression strategies over an eight-week lower-body resistance-training cycle. One group progressed by increasing load while keeping repetitions in an 8–12 repetition range. The other group kept the initial load and progressed repetitions. The study used trained young adults, lower-body exercises, two sessions per week, four sets per exercise, and sets performed to momentary concentric failure. This is a useful study because it asks whether the athlete must always add weight to keep progressing.

The answer was no. Both groups improved. Muscle-size changes were broadly similar in most measured regions. The load-progression group showed a small, uncertain advantage on the squat strength outcome, while one hypertrophy outcome modestly favoured the repetition-progression group. The intervals crossed practically unimportant values and, depending on the outcome, included zero. The conclusion was that both load and repetition progression appeared viable over the eight-week cycle.

The limitation is exactly the point for this project: Plotkin did not randomise athletes to 2.5%, 5%, and 10% increments. The load group was told to increase load while staying in the target range, but the paper does not turn that into a universal increment formula. The study also does not isolate load-increment size from the broader choice of progression axis. A repetition-progression group can accumulate a different external volume-load than a load-progression group; that is part of the intervention, not a nuisance that disappears.

Plotkin therefore supports a dual-axis engine. If a load jump is too large, the athlete can still make legitimate progress through repetitions, execution quality, reduced RPE at the same load, or another stated lever. The paper supports the existence of a pressure-release valve. It does not support the exact pressure setting.

This distinction protects the product from a common implementation mistake: if the engine treats a successful 10-repetition set as proof that the next load must increase, it will force a plate jump even when the equipment jump is disproportionate. A better engine says: “The athlete has demonstrated progress. The next available load jump is 10% and exceeds this movement’s configured jump ceiling. Hold the load, progress the repetition target or confirm with another comparable exposure.” That behaviour is consistent with Plotkin’s result and honest about what the paper did not test.

## 5. Direct evidence on increment size

The direct evidence is sparse. Hostler et al. studied the effectiveness of very small increments in two upper-body exercises, the bench press and triceps press. The small-increment protocol used changes of approximately 0.22 kg after completing seven to eight repetitions and approximately 0.44 kg after completing nine or more. The traditional comparison used substantially larger exercise-specific increments. Both approaches improved strength over eight weeks. The authors concluded that small-increment progression appeared as effective in the short term, while preliminary extended observations suggested that the traditional group might eventually increase resistance more effectively.

This study matters because it defeats the assumption that a small increment is automatically too small to work. It does not prove that the smallest possible increment is best. It is small, upper-body-specific, short, and not a clean answer to adherence or long-term stalling. It also does not tell us how to compare a 0.5 kg increase on a 25 kg press with a 5 kg increase on a 180 kg deadlift.

Buskard et al. compared four approaches to load progression in healthy older adults across an eleven-week structured resistance-training period. The groups differed in how load increases were determined, including repetition maximum, percentage-based, RPE, and repetitions-in-reserve approaches. No significant between-group differences emerged for the strength or functional-performance outcomes. The RPE group found the training more tolerable and enjoyable. This is not a direct increment-size trial, but it supports the principle that multiple reasonable progression policies can produce similar outcomes and that tolerability is not a trivial product variable.

The direct answer to the project’s main question is therefore:

> No convincing experiment was located that randomised comparable trainees to approximately 2.5%, 5%, and 10% progression increments while holding exercise, target repetitions, progression trigger, volume, frequency, and context constant.

That negative finding is load-bearing. It means the final number must be treated as a tunable starting parameter. The correct response is not to give up and use a flat 2.5 kg rule. It is to make the percentage visible, constrain it by equipment, provide a repetition fallback, and collect data that can later test the engine’s own behaviour.

## 6. Why relative progression is the least-wrong default

A flat 2.5 kg change has radically different meaning at different loads. At 25 kg it is a 10% increase. At 50 kg it is 5%. At 100 kg it is 2.5%. At 180 kg it is about 1.4%. Therefore a global “add 2.5 kg when successful” rule silently assigns an aggressive progression to light exercises and a conservative progression to heavy ones.

A percentage rule is not physiologically perfect. A 5% external-load increase is not guaranteed to create a 5% increase in internal stress because leverage, range of motion, technique, machine geometry, fatigue, and bodyweight alter the task. Nonetheless, relative loading is a more coherent control variable than a fixed kilogram constant when one engine must serve different exercises and users.

The correct algorithm is a two-stage policy. First, calculate a desired relative change. Second, ask whether the available equipment can express that change without creating an excessive actual jump. The engine must store both values:

```text
target_delta = anchor_load × progression_percent
actual_delta = selected_load − anchor_load
actual_jump_percent = actual_delta / anchor_load
```

If the smallest available increment at or above the target creates an actual jump below the movement’s ceiling, use it. If it exceeds the ceiling, do not pretend that the result equals the target. Hold the load and use a permitted fallback: add a repetition, improve the achieved RPE margin, add a set only if the programme allows that lever, or schedule a later confirmation exposure. The fallback is not a failure to progress. It is the correct response to coarse hardware.

The engine should not always choose the smallest available increment below the target either. A 0.5% increase may be too small to matter or may be unavailable. The selection policy should be deterministic and configurable. A safe default is to choose the smallest available jump that reaches the desired target without exceeding the exercise cap. When no such jump exists, hold and progress another axis. For a stable lower-body compound with 180 kg as the anchor, a 2.5 kg jump is a 1.39% increase and is acceptable even if it is below a 2.5% target. For a 25 kg upper-body movement, a 2.5 kg jump is 10%; the engine should hold or microload rather than force it.

The cap is also a heuristic, not a scientific law. The recommended initial values are 5% for small-load or upper-body movements and 7.5% for large lower-body movements when equipment leaves no better choice. Those caps should be versioned and tested. They are a safety-oriented translation of the ACSM band and plate realities, not a claim about optimal adaptation.

## 7. The progression controller

The controller should not make a decision from a single “success” boolean. A set can be completed because the athlete overshot the intended RPE, because the target was too easy, because the exercise was changed, or because the user entered a value without actually performing the set. The engine needs a validity layer before a progression layer.

### 7.1 Exposure validity

An exposure is comparable only if the exercise identity, meaningful range of motion, target mode, load units, set structure, and general technique are sufficiently stable. A changed machine, exercise substitution, unusual pain, or a major change in rest interval may still be useful data, but it should not automatically count as equivalent evidence for a load jump.

Each exposure should record:

- exercise and variation identity;
- intended load and actual load;
- target repetitions or time;
- completed repetitions or time;
- set count and order;
- reported RPE or RIR;
- technique or range-of-motion flag;
- pain/symptom flag;
- rest interval where it matters;
- preceding conditioning load and recent training gap;
- data completeness and source;
- whether the athlete used an approved substitution.

### 7.2 Success classification

A valid success is not merely “all boxes ticked.” The first implementation should classify an exposure as `successful`, `successful_but_uncertain`, `held`, `missed`, `pain_blocked`, `invalid_for_progression`, or `incomplete`. A successful exposure may be uncertain if RPE is missing, technique changed, equipment changed, or conditioning fatigue was unusually high. A valid completion under a lower dose can be a successful recovery-mode exposure without authorising progression.

### 7.3 Promotion gate

The base gate is two comparable successful exposures, because repeated confirmation is a practical noise filter and aligns with the broad repeated-performance logic in ACSM. The number is explicitly a product heuristic. The engine should allow the configuration to change, and telemetry should record what would have happened under one, two, or three exposures during validation replay.

When the gate is met, promote only one lever. Do not increase load, repetitions, set count, and conditioning intensity simultaneously. One-lever changes preserve causal clarity and reduce the risk that a temporary good day becomes an overlarge training dose.

### 7.4 Hold state

The hold state is a first-class outcome, not a failure. Use it when the athlete completes the work but the evidence is too noisy to justify a change; when the next load jump is too coarse; when readiness and performance conflict; or when missing information matters to the decision. The user-facing explanation should say what was preserved and what evidence is missing.

### 7.5 Repetition fallback

When load cannot rise safely, use repetition progression only inside the exercise’s prescribed range and fatigue budget. Do not turn a strength-focused set into an endless AMRAP merely because the plates are coarse. A fallback can be as small as one additional repetition on the next exposure, a reduced RPE target at the same load, or a second confirmation at the same load. The engine must retain the session purpose.

## 8. RPE and RIR: useful instruments, not truth oracles

RPE/RIR is valuable because a percentage based on an outdated maximum can be wrong on a particular day, while the athlete’s actual effort and completion can provide immediate information. Studies in trained lifters show meaningful relationships between perceived effort, velocity, and proximity to failure. But perception has error. Users are less accurate when many repetitions remain, the exercise is unfamiliar, the set is technically unstable, or the athlete has not been familiarised with the scale.

The engine should therefore use RPE/RIR conditionally. For familiar exercises in trained users and sets close enough to failure for the scale to be meaningful, it can inform within-session adjustment and progression confidence. For a novel exercise, a high-repetition set, pain-affected movement, or an RPE entry that conflicts with the observed repetitions and technique, it should be advisory only.

The app should learn user-specific calibration rather than assuming that every athlete’s RPE 8 means the same thing. Calibration does not require a complicated machine-learning model. The engine can calculate simple bias summaries: how often an athlete’s reported RPE 8 was followed by a large overshoot, how often an RPE 9 ended early, and how stable the relationship is for a specific exercise. These summaries should be shown as evidence quality, not as a permanent personality label.

RPE/RIR should not be used to make the claim that autoregulation is always superior to fixed percentages. Trials and reviews are mixed. In some comparisons, the autoregulated group achieved a different actual intensity, which makes it difficult to isolate the decision method. Other studies found similar gains. The correct product claim is that RPE/RIR is an additional input that can help the engine remain responsive when a fixed percentage is stale.

Within-session correction must also be separated from cross-session programming. If the athlete misses at 100 kg and the session corrects to 94 kg, 94 kg is the effective load for that session. It is not automatically the new long-term anchor. The engine should retain `session_opening_load`, `effective_load`, and `last_successful_anchor_load` separately. Otherwise a failed session can cause an invisible compound reduction when the next cross-session deload is calculated from an already reduced number.

## 9. Deloads: common, plausible, and under-tested

Deloading is one of the clearest areas where coaching practice outruns direct evidence. A 2024 survey of 246 competitive strength and physique athletes found that deloads were commonly used for about 6.4 days every 5.6 weeks. Athletes typically reduced volume through sets and repetitions, often reduced load and effort, and generally kept exercise selection and frequency similar. This describes what experienced athletes report doing. It does not prove that every athlete benefits from a deload on that schedule.

The 2023 Delphi consensus provides useful terminology and a practical framework, but consensus is not a dose-response trial. The 2024 Coleman study found that a one-week midpoint deload involving training cessation could negatively affect some lower-body strength outcomes compared with continuous training, even when hypertrophy outcomes were similar. The 2026 Pancar within-subject study in 19 untrained young men found that planned reductions in weekly sets and frequency at weeks 4 and 8 did not hinder hypertrophy or 10RM strength-endurance over eight weeks. These studies are not contradictory once their protocols are respected: they tested different populations, goals, exercise selections, deload definitions, and outcome measures.

The lesson is not “deloads are bad” or “deloads are mandatory.” The lesson is that deload is a family of interventions. Reducing sets while keeping some exposure is not the same as complete cessation. A reactive deload after corroborated decline is not the same as a pre-planned week inserted on a calendar. A strength test is not the same as hypertrophy or local endurance. The app must preserve those distinctions.

The final reactive rule is:

1. One comparable miss: hold, inspect context, and do not automatically deload.
2. Repeated comparable misses or corroborated deterioration: enter a conservative reduction state.
3. Isolated movement failure with normal session performance: reduce the affected load or repeat the exposure; keep overall volume unless there is reason not to.
4. Several related movements failing, or performance declining across the session: reduce one principal stress lever, commonly volume or intensity, and reassess.
5. System-wide poor performance with high conditioning fatigue, illness, or poor readiness: reduce volume and/or intensity according to the configured safety policy, but do not stack unexplained reductions.

If a default number is required for the first release, use 5% from the last successful opening anchor. Label it `product_heuristic`. Escalate toward 7.5–10% only when broader evidence supports it, and consider reducing volume rather than merely lowering load when the problem appears systemic. Never calculate a nominal “5% deload” from a load that has already been walked down by a within-session miss unless the product explicitly wants the compound effect and displays it.

## 10. Why “two consecutive misses” is not science—but can still be useful

No direct trial was located that compares one, two, and three consecutive failures as trigger thresholds for load reduction. The threshold is therefore not physiology. It is a control-system debounce choice. A single missed set contains many possible explanations: sleep, warm-up, technique, rest, equipment, illness, work stress, preceding conditioning, or a bad RPE estimate. Requiring repeated comparable evidence reduces the chance that the engine overreacts to noise.

The word “comparable” carries more weight than the number two. Two failures after a heavy conditioning session and a poor night’s sleep are not the same evidence as two failures under normal conditions. Two misses on a different machine are not the same as two misses on the same exercise. A failed set because of pain should not enter the same counter as a non-painful strength miss. A missing set should not be treated as a completed failure.

The engine should maintain separate counters or evidence states:

- `performance_miss_streak` for comparable non-painful misses;
- `pain_block_state` for symptom-related interruption;
- `systemic_decline_evidence` for broad cross-movement deterioration;
- `data_quality_conflict` for contradictory or incomplete records;
- `calibration_state` for a return after a meaningful gap.

This is superior to a single “failed sessions” integer because it prevents unrelated events from being combined into a fake pattern. If the product UI wants to show “two misses,” it should be able to explain which two exposures counted and why.

## 11. Hybrid training changes the meaning of a miss

A strength-only programme can interpret a missed squat mainly through the recent strength plan. A hybrid athlete may have performed intervals, long Zone 2, a physically demanding job, poor sleep, or another lower-body session nearby. The same failed squat may represent local strength limitation, systemic fatigue, glycogen availability, dehydration, accumulated conditioning load, or a technical issue.

The engine should not pretend that a single composite readiness score solves this. It should preserve modality-specific evidence and use a transparent priority order. Safety and pain come first. Actual performance and technique come next. Session purpose follows. Recent comparable strength evidence matters. Wearable or HRV data is advisory context.

The programming layer should store conditioning exposure in a way that permits interpretation: modality, duration, intensity, heart-rate zones, interval count, perceived effort, and time since the session. It should not silently convert a hard row into “leg fatigue” or use a cycling metric as though it were a squat readiness measure. The system can infer that a hard lower-body interval session may be relevant context, but it must label the inference.

The most useful hybrid behaviour is not to cancel every strength session after a hard conditioning day. It is to preserve the primary purpose with a controlled adjustment when multiple signals converge. A normal strength performance despite a low wearable recovery score should not be demoted automatically. Poor performance plus poor perceived recovery plus unusually high preceding conditioning load is stronger evidence for reducing the day’s dose. This multi-signal logic remains an engineering policy; no reviewed study validates the exact fusion algorithm for this product.

## 12. HRV and wearables: advisory context only

HRV-guided training has more support in endurance distribution than in strength-load prescription. Individualised endurance studies show that using HRV can alter the distribution of hard and easy training and sometimes produce useful outcomes. Strength studies are smaller and mixed. HRV is influenced by measurement protocol, sleep, illness, stress, posture, breathing, alcohol, and device quality. It is not a direct measure of local muscle capacity, tendon status, pain, technical readiness, or medical safety.

The product rule is therefore simple: HRV can lower confidence or support a conservative option when it agrees with performance and context, but HRV alone must not increase load, authorise high-risk work, diagnose overtraining, or override pain. The engine should normalise HRV within the athlete and measurement protocol, retain the raw value and timestamp, and mark stale or missing data. Cross-device values must not be treated as directly interchangeable without a declared normalisation policy.

The UI should avoid an authoritative “recovery score says train hard” message. A better statement is: “Today’s wearable recovery signal is lower than your usual range. Your recent comparable performance is normal and no pain was reported, so the planned session remains available, but progression is held until performance provides stronger evidence.” That statement tells the truth about both the signal and its limits.

## 13. Return after a layoff

Detraining evidence does not provide a universal time-off-to-load-reduction equation. Short breaks can preserve much of the strength and muscle adaptation, but performance expression, exercise skill, conditioning, connective tissue tolerance, and confidence may change at different rates. A trained powerlifter returning after two weeks is not the same case as a beginner returning after three months, and a barbell lift is not the same as a running interval.

The engine should represent a return as calibration. It should reduce complexity, avoid unnecessary failure exposure, and observe current technique, RPE/RIR, completion, and symptoms. The first session should gather information without treating the athlete as untrained or forcing a fixed percentage reduction. The exact amount of load or volume reduction can be configured by modality and risk profile, but the product documentation must state that it is a conservative policy rather than a validated equation.

The calibration state should expire based on successful comparable exposures, not merely time. A user who returns and completes two stable exposures with normal technique may leave calibration sooner than a user who returns with pain, repeated misses, or contradictory data. The system must not silently promote a calibration session to a normal progression anchor without recording that it was a calibration exposure.

## 14. Pain is not fatigue

Pain-monitoring evidence can support continued activity in specific rehabilitation contexts. Silbernagel’s Achilles tendinopathy work is an important example of a condition-specific protocol with pain and next-morning monitoring. It cannot be converted into a universal “5/10 pain is safe” rule for every user and exercise. A consumer training app usually does not know the diagnosis, tissue state, red flags, or rehabilitation plan.

Pain, exertional fatigue, breathlessness, soreness, technical breakdown, and fear are different signals. The data model must keep them separate. A pain event should stop, modify, or substitute the affected movement according to safety policy, record location and quality, and block automatic progression for that movement or pattern until the symptom state is updated. The engine should offer a safe alternative where appropriate, but it must not diagnose or reassure beyond its evidence.

The human-logic behaviour is bounded: distinguish a difficult set from a painful set; preserve the session purpose through a safe substitution or reduced dose when possible; do not let motivation override a hard safety state; and make the next check explicit. This is a product safety policy informed by clinical evidence boundaries, not a universal treatment protocol.

## 15. Trends, reliability, and the minimum data problem

No universal number of sessions was found that makes an e1RM or volume trend trustworthy. The correct number depends on measurement error, exercise consistency, sampling frequency, missingness, and the magnitude of change that matters to the decision. Three comparable exposures may be enough to decide “hold and gather more data” but not enough to claim a long-term trend. Ten non-comparable exposures may be less informative than three standardised ones.

The engine should use reliability concepts rather than a magical minimum. Store the raw observations, standardise what can be standardised, estimate within-user variability when enough repeated data exists, and compare observed change with the estimated error band. Typical error, standard error of measurement, limits of agreement, and minimum detectable change are more informative than a raw correlation alone.

For e1RM, store the formula version, load, repetitions, RPE/RIR, exercise, technique flags, and whether the set was valid for trend use. Do not mix a rep-max estimate from a machine with a barbell estimate without an explicit modality boundary. Smoothing can reduce noise, but a three-point moving average is not a universal law. It is a declared filter. The engine should show how many observations were used and whether the window contains gaps or exercise changes.

The practical product states are:

- `insufficient_evidence`: not enough comparable data;
- `uncertain_change`: observed change does not exceed estimated noise;
- `stable`: no meaningful change detected;
- `probable_improvement`: change exceeds current error estimate with adequate data;
- `probable_decline`: change exceeds current error estimate and is corroborated;
- `conflicted`: signals disagree or protocol changed.

This is more honest than showing “progress +4.3%” without explaining whether the number is a stable trend or one unusually good set.

## 16. Confidence and human-like logic

A confidence score is meaningful only if it is confidence in a defined decision over a defined horizon. “Confidence 72%” is not useful if the user does not know whether it means the load will be completed, that recovery is adequate, or that hypertrophy will occur. Calibration research also shows that an apparently confident estimate can be poorly matched to real-world correctness.

The engine should therefore communicate evidence state rather than decorative precision. Use statuses such as `approved`, `held`, `reduced`, `repeated`, `substituted`, `calibration`, `blocked`, `needs_input`, and `needs_clinical_review`. Each decision should expose observed signals, estimated state, conflicts, missing data, reason codes, evidence status, and the next recheck point.

The human-logic layer is not an unconstrained language model. It is a bounded policy that behaves like a calm coach: preserve the session purpose, distinguish danger from inconvenience, avoid punishment and make-up debt, and choose the smallest safe change. An AI language layer may translate the deterministic decision into natural language later, but it must not change the decision or invent a rationale.

The priority order is safety gate, pain/symptoms, immediate training quality, session purpose, real-world constraints, recent comparable evidence, and advisory wearable signals. Lower-priority context cannot override a higher-priority safety state. A good readiness score cannot override pain. A low HRV score cannot override strong performance and safe technique by forcing a demotion. A user request to “push through” cannot override a blocked state.

## 17. Final implementation specification

The following is the locked first-release specification. Numbers marked as heuristics must be versioned and exposed in the engine metadata.

### 17.1 Core constants

```text
DEFAULT_PROGRESS_TARGET = 0.025       # product heuristic inside ACSM 2–10% band
SMALL_MOVEMENT_JUMP_CAP = 0.05        # product heuristic
LARGE_MOVEMENT_JUMP_CAP = 0.075       # product heuristic
DEFAULT_REACTIVE_DELOAD = 0.05        # product heuristic
ESCALATED_DELOAD_RANGE = 0.075–0.10   # only with corroborated systemic decline
SUCCESS_CONFIRMATION_EXPOSURES = 2    # product heuristic / debounce rule
SINGLE_MISS_ACTION = hold_and_context  # product policy
PAIN_ACTION = block_or_modify          # safety policy
LAYOFF_ACTION = calibration             # product policy
HRV_ROLE = advisory                     # evidence boundary
```

### 17.2 Anchor and load fields

```ts
type ExerciseProgressState = {
  exerciseId: string;
  variationId: string;
  lastSuccessfulOpeningLoad: number | null;
  currentOpeningLoad: number | null;
  currentEffectiveLoad: number | null;
  targetIncrementPercent: number;
  actualIncrementPercent: number | null;
  availableLoadStep: number | null;
  successStreak: number;
  comparableMissEvidence: number;
  painBlock: boolean;
  calibration: boolean;
  lastDecision: DecisionStatus;
  engineVersion: string;
};
```

`lastSuccessfulOpeningLoad` is the cross-session anchor. `currentOpeningLoad` is what the user was asked to start with today. `currentEffectiveLoad` is what was actually used after within-session corrections. These must not be collapsed into one field.

### 17.3 Promotion pseudocode

```text
if safety_gate or pain_block:
    return blocked_or_modified

if calibration:
    return calibration_prescription

if not comparable_success:
    if valid_nonpainful_miss:
        record miss evidence
        return hold
    return insufficient_evidence

if success_streak < confirmation_exposures:
    return repeat_current_prescription

target_delta = anchor_load × target_increment_percent
candidate = choose_available_load(anchor_load, target_delta)
actual_jump = (candidate - anchor_load) / anchor_load

if candidate exists and actual_jump <= movement_jump_cap:
    promote load only
else:
    hold load and activate permitted repetition/RPE fallback
```

### 17.4 Reactive-reduction pseudocode

```text
if pain or red_flag:
    do not calculate a normal deload
    enter safety pathway

if one comparable nonpainful miss:
    hold from last successful anchor
    collect context

if repeated comparable misses and local problem only:
    new_opening = round(anchor_load × (1 - 0.05))
    keep volume unless other evidence says otherwise

if broad decline across related movements:
    choose one primary reduction lever
    prefer volume reduction when fatigue is systemic
    use 0.05–0.075 load reduction if load reduction is selected

if systemic decline, illness, or extreme conditioning stress:
    use conservative mode or recovery mode
    do not stack hidden reductions
```

### 17.5 Explanation contract

```text
What I noticed: [observed facts].
What that means today: [bounded interpretation and confidence].
What we are doing: [specific action and preserved session purpose].
Why: [reason codes in plain language].
What would change it: [next check or missing information].
```

The explanation is generated from structured fields. It is not allowed to invent physiology, diagnose fatigue, or claim that a heuristic is a study result.

## 18. Worked examples

### Example A: 25 kg upper-body movement with coarse plates

The athlete’s last successful opening load is 25 kg. The target increment is 2.5%, or 0.625 kg. The gym can only express a 2.5 kg increase, which would be 10%. The movement cap is 5%. The correct result is not 27.5 kg disguised as a 2.5% progression. The engine holds 25 kg, advances the repetition or RPE-quality target within the planned range, and records that equipment resolution prevented the target load change.

### Example B: 60 kg lower-body movement

The target increment is 1.5 kg. A 2.5 kg step creates a 4.17% jump, below the configured large-movement cap. The engine can select 62.5 kg, provided the exposure was comparable and the success gate was met. It records target 2.5%, actual 4.17%, and equipment rounding.

### Example C: 100 kg movement with a missed set

The session opens at 100 kg. The athlete misses the target, and the within-session rule corrects the effective load to 94 kg. The session record stores opening 100, effective 94, and failed exposure. If repeated comparable misses later trigger a 5% reactive reduction, the next opening load is calculated from the last successful anchor, 100 kg, producing 95 kg before equipment rounding. It is not calculated from 94 kg unless the product explicitly chooses a compounded policy and displays it.

### Example D: low HRV but normal performance

Wearable recovery is lower than the athlete’s usual range. The athlete completes the planned warm-up and working sets with normal technique, target repetitions, and expected RPE. The engine does not force a demotion or progression. It completes the planned purpose and holds progression if the low HRV creates uncertainty. The explanation says the wearable signal was advisory and performance was the stronger direct evidence today.

### Example E: pain during a press

The athlete reports sharp shoulder pain during the third set. The engine does not treat this as RPE 10. It enters a pain pathway, stops or modifies the movement, records location and symptom detail, and offers a safe substitution only within the configured policy. Automatic progression for the affected movement is blocked pending symptom update. The app does not diagnose the shoulder.

### Example F: missed training week

The athlete returns after seven days away. The engine enters calibration for the next exposure, caps effort, and observes current performance. It does not automatically apply a universal 10% reduction. If the user completes the calibration exposure with stable technique and no symptoms, the next session can return toward normal prescription with the calibration evidence recorded.

## 19. Validation plan for THE Hybrid Engine itself

Because the literature does not identify the exact constants, the app must be designed to learn whether its own policy is useful. This is not permission to claim that a single-user log proves general efficacy. It is a product validation programme.

### Phase 1: deterministic replay

Build fixtures in which identical inputs produce identical outputs across browser, coach, offline, and cloud-rehydrated execution. Include every decision state and reason code. A changed engine version must be able to explain why its output changed.

### Phase 2: boundary testing

Test loads around the jump cap: 24.9, 25, 25.1, 49.9, 50, 50.1, 99.9, 100, 100.1, and heavy lower-body anchors. Test available increments below, equal to, and above the target. Verify that displayed target percentage and actual percentage cannot be confused.

### Phase 3: evidence-quality testing

Replay missing RPE, missing HRV, missing heart rate, changed exercise, changed equipment, incomplete session, pain, illness flag, unusual conditioning, poor sleep, and conflicting user/device reports. The expected behaviour is not always to stop. It is to lower confidence, hold, reduce, substitute, or request the smallest useful clarification according to the state.

### Phase 4: historical replay

Replay real completed sessions without changing the historical record. Compare the current rule, a flat-kilogram rule, a 2.5% rule, a 5% rule, and a repetition-fallback rule. Evaluate missed-session frequency, unnecessary load jumps, repeated stalls, total completed work, pain-block violations, and the frequency of “hold” decisions. The goal is not to crown a single metric. It is to see whether the engine behaves proportionately.

### Phase 5: prospective pilot

Run the locked policy prospectively with the user or a small invited group. Log decisions and overrides. Measure whether users understand the reason for a hold, whether they follow safe substitutions, whether they perceive the system as useful, and whether the engine’s decisions remain stable under work stress and hybrid scheduling. Any outcome claims must remain exploratory until a properly designed study exists.

### Primary product metrics

- decision determinism;
- percentage of adaptive decisions with complete reason codes;
- rate of silent data coercion, which must be zero;
- rate of progression decisions with a valid comparable exposure;
- rate of pain events that incorrectly enter normal fatigue logic, which must be zero;
- unnecessary progression after a single noisy observation;
- user comprehension of hold and calibration states;
- user override rate and whether overrides occur in blocked states;
- recovery of incomplete sessions after reload or offline use;
- integrity of coach-target versus athlete-result ownership.

## 20. Product architecture required by the evidence

The evidence conclusion is inseparable from the project’s existing architecture. The athlete and coach apps must remain separate surfaces connected by one tested data boundary. The coach owns targets, prescriptions, exercise instructions, and planned structure. The athlete owns actual load, actual repetitions, RPE, notes, completion, symptoms, and session results. The engine may calculate an adaptive recommendation, but that recommendation must be a versioned target snapshot, not a mutation of historical athlete results.

The application must remain local-first and work offline. Local storage is the first write path; cloud sync is optional, debounced, and reconciled by stable identifiers. Authenticated secrets and WHOOP provider credentials remain server-side. The adaptive layer must not become an excuse to collapse the data ownership boundary or create a hidden server-only brain that the user cannot audit.

The database and sync model should preserve raw data and derived state separately. Raw logs are facts: what the athlete entered or what the device supplied. Derived fields are estimates: e1RM, trend, fatigue state, confidence, or recommendation. If the algorithm changes, the raw facts should remain available for replay. This is the technical equivalent of the evidence rule that a product heuristic must remain testable.

## 21. Claims the product may and may not make

The product may say that it uses a conservative, performance-informed, equipment-aware progression policy. It may say that the progression target is inspired by ACSM’s 2–10% operating range. It may say that load and repetition progression are both supported pathways. It may say that the system uses RPE/RIR conditionally, wearable data as advisory context, and explicit hold states when evidence is insufficient.

The product must not say that 2.5% is the scientifically optimal increase. It must not say that two missed sessions is a validated deload threshold. It must not say that a 5% reduction is proven to restore performance. It must not say that HRV is a direct strength-readiness or injury-readiness score. It must not say that a universal pain score is safe. It must not imply that MacroFactor, JuggernautAI, or RTS independently validated this product. It must not imply that a confidence number predicts hypertrophy or injury prevention unless that exact outcome has been calibrated and tested.

## 22. Final answers to the original questions

### Does ACSM actually recommend a percentage increase?

Yes. The 2009 position stand recommends a 2–10% increase after the athlete exceeds the desired workload by one to two repetitions on two consecutive sessions, with smaller increases for smaller exercises and larger increases for larger exercises. The recommendation is real and useful. Its exact dose is not established by a direct percentage-comparison trial.

### What did Plotkin test?

Plotkin compared progressing load while staying within a repetition range with progressing repetitions at the starting load over eight weeks of lower-body training in trained young adults. Both strategies produced meaningful adaptations. It did not test 2.5%, 5%, or 10% increments and cannot justify a particular percentage constant.

### Does a trial compare increment size directly?

No convincing trial was found that cleanly compares approximately 2.5%, 5%, and 10% progression increments under otherwise matched conditions. Hostler supports the feasibility of microloading in upper-body exercises. Buskard supports the practical viability of multiple progression methods. Neither identifies a universal optimum.

### Is two consecutive failure the right deload threshold?

There is no direct evidence that two is superior to one or three. It is a defensible product debounce rule only when “comparable” is defined and pain, illness, missing data, and unusual conditioning are kept separate. The engine should treat corroborated deterioration as more important than the raw count.

### Should deload reduce load or volume?

The evidence does not establish one universal lever. Surveys suggest athletes commonly reduce volume, effort, and sometimes load. A narrow 2026 study found that planned reductions in sets and frequency did not hinder hypertrophy or 10RM strength-endurance in untrained young men. For the engine, isolated local failure can justify a load hold or reduction; broader decline should often reduce volume or overall stress rather than merely lowering the bar weight.

### What is the final implementation?

Use a 2.5% target from the last stable opening load, equipment-aware rounding, movement-specific jump caps, and a repetition/RPE fallback. Promote after two comparable successful exposures as a product heuristic. Hold after one miss. After repeated comparable deterioration, reduce from the last successful anchor by a default 5% and escalate only with corroborating systemic evidence. Keep pain, HRV, layoff, conditioning, and data quality in separate state paths. Record every decision with reason codes and evidence status.

## 23. Closure statement

This research question is now closed for the first product build. Further literature searching may find another adjacent study, but it is unlikely to change the central conclusion without a direct increment-size or reactive-deload trial. The engine does not need another round of searching before implementation. It needs the rules above encoded, tested, and measured.

The project should proceed with the following sentence in its code documentation:

> THE Hybrid Engine uses a versioned, deterministic, evidence-informed progression and fatigue policy. Its 2.5% progression target and 5% reactive reduction are configurable product heuristics selected inside broader evidence-supported practice ranges; they are not presented as experimentally established optima. The engine prioritises safety, pain separation, comparable performance, equipment-aware rounding, explicit uncertainty, and preservation of the session purpose.

That is the end of the research loop. Build the system, test the system, and let future user data challenge the heuristics rather than hiding them behind the word “science.”

## Appendix 1. Evidence ledger for the load-progression decision

The following ledger records the major sources used in the final judgement. It is intentionally explicit about transferability.

| Source | Design / population | What it contributes | What it cannot prove | Final status |
|---|---|---|---|---|
| Ratamess et al., 2009, ACSM position stand, DOI 10.1249/MSS.0b013e3181915670 | Position stand for healthy adults | 2–10% operating band after repeated performance; exercise-size distinction | Exact optimum; plate algorithm; deload rule | Position-stand guidance, Category B |
| Kraemer et al., 2002, ACSM position stand, PMID 11828249 | Earlier position stand | Historical continuity of repeated-performance progression logic | Increment-size dose response | Position-stand guidance |
| Feigenbaum & Pollock, 1999, PMID 9927008 | Prescription/review paper | Source lineage behind practical progression language | Direct comparison of 2.5/5/10% | Prescription precedent |
| Plotkin et al., 2022, PMID 36199287 / PMCID PMC9528903 | Randomised load versus repetition progression; trained young adults | Load and repetition progression both viable over eight weeks | Increment-size optimum; hybrid generalisation | Direct but narrow |
| Hostler et al., 2001, PMID 11708713 | Small-increment upper-body trial | Microloading can work over eight weeks | Universal long-term optimum; lower-body transfer | Direct but small |
| Buskard et al., 2019, DOI 10.1249/MSS.0000000000002038 | 82 older adults; four progression methods | No clear strength/functional winner; RPE tolerability signal | Young hybrid athlete; exact increments | Direct adjacent |
| Teixeira et al., 2019, PMCID PMC6616272 | Novel load-progression strategy | Shows progression can vary across load/volume and still be studied | Universal practical prescription | Adjacent |
| Schoenfeld et al., 2021, PMCID PMC7927075 | Loading recommendations review | Goal-specific load and rep interpretation | Increment-size optimum | Review / adjacent |
| Huang et al., 2025, PMCID PMC12336695 | Autoregulation network meta-analysis | Suggests autoregulated methods may improve strength in pooled data | Product-specific rule; clean separation of achieved intensity | Emerging pooled evidence |
| Hickmott et al., 2022, PMCID PMC8762534 | Autoregulation systematic review | Comparative evidence is mixed and heterogeneous | Universal RPE superiority | Systematic review |
| Helms et al., 2018, DOI 10.3389/fphys.2018.00247 | RIR versus percentage-style training | No clear universal RIR superiority in small trial | All populations/exercises | Direct but small |
| Graham & Cleather, 2021, DOI 10.1519/JSC.0000000000003164 | RIR / autoregulation comparison | Supports possible strength benefit with different achieved intensity | Exact product controller | Direct but confounded by achieved dose |
| Zourdos et al., 2016, DOI 10.1519/JSC.0000000000001049 | RPE/velocity relation | Supports RPE as a structured measurement | Perfect accuracy | Direct acute evidence |
| Helms et al., 2017, DOI 10.1519/JSC.0000000000001517 | RPE/RIR reliability in trained lifters | Supports familiarisation and conditional use | Universal calibration | Direct small study |
| Mansfield et al., 2020, DOI 10.1519/JSC.0000000000003779 | RIR accuracy | Estimates can be wrong, especially away from failure | One universal error correction | Adjacent |
| Steele et al., 2017, DOI 10.7717/peerj.4105 | RIR and failure interpretation | Highlights conceptual and measurement issues | App-ready threshold | Conceptual/adjacent |
| Vieira et al., 2022, PMID 34881412 | Failure-training systematic review/meta-analysis | Failure adds acute fatigue and performance cost | Two-miss deload threshold | Adjacent pooled evidence |
| Rogerson et al., 2024, DOI 10.1186/s40798-024-00691-y | Survey of 246 competitive athletes | Describes common deload timing and levers | Optimal causal schedule | Survey / practice description |
| Bell et al., 2023, PMCID PMC10511399 | International Delphi consensus | Defines deloading terminology and practice | Efficacy or dose-response | Consensus |
| Coleman et al., 2024, PMCID PMC10809978 | One-week midpoint deload comparison | Shows planned cessation can affect strength outcomes differently from hypertrophy | Reactive deload rule | Direct narrow study |
| Pancar et al., 2026, DOI 10.1038/s41598-026-40612-5 | Randomised within-subject; 19 untrained men | Reduced sets/frequency did not hinder 8-week hypertrophy/10RM | Trained hybrid athletes; best deload dose | Direct narrow study |
| Ogasawara et al., 2013, PMID 23053130 | Periodic training with 3-week cessation | Long breaks did not erase all adaptations in a narrow protocol | Normal 7-day reactive deload | Direct narrow study |
| Bickel et al., 2011 | Reduced maintenance volume | Small volume may maintain adaptation | Reactive deload efficacy | Adjacent |
| Tavares et al. | Reduced volume/frequency after training | Reduced training can maintain prior gains in some contexts | Universal return formula | Adjacent |
| Atkinson & Nevill, 1998, PMID 9820922 | Reliability methods review | Error, limits of agreement, and repeatability principles | Universal session count | Methods review |
| Hopkins, 2000, PMID 10907753 | Reliability and validity methods | Typical error and meaningful-change logic | Product-specific window | Methods review |
| Grgic et al., 2020, PMID 32681399 | 1RM reliability systematic review | Reliability depends on test context and protocol | Universal e1RM confidence | Systematic review |
| Bellenger et al., 2016, DOI 10.1007/s40279-016-0484-2 | HRV-guided endurance review/meta-analysis | Possible benefit in endurance distribution | Strength-load gate | Systematic review |
| Manresa-Rocamora et al., 2021, DOI 10.3390/ijerph181910299 | Individualised HRV endurance study | Supports personalisation but not universal cutoff | Mixed-modality app rule | Direct adjacent |
| de Oliveira et al., 2019, DOI 10.1080/17461391.2019.1572227 | HRV-guided resistance study | Small and limited evidence | HRV-only superiority | Direct small |
| Bittencourt et al., 2024, DOI 10.3389/fphys.2024.1472702 | HRV-individualised recovery in older women | Tests autonomic-guided recovery | Young hybrid generalisation | Direct narrow |
| Thamm et al., 2019, DOI 10.3390/ijerph16224353 | HRV and resistance-exercise response | HRV not a direct local muscle measure | Tissue readiness | Adjacent |
| Hwang et al., 2017, PMID 28328712 | Short detraining in trained men | Retention after a short break can be substantial | Fixed return percentage | Direct narrow |
| Mujika & Padilla, 2000, PMID 10966148 | Detraining review | Modality- and duration-specific detraining | Universal strength equation | Review |
| Encarnação et al., 2022, DOI via MDPI review | Detraining review | Strength and hypertrophy time courses vary | App-ready percentage | Review |
| Silbernagel et al., 2007, DOI 10.1177/0363546506298279 | Achilles pain-monitoring RCT | Condition-specific symptom monitoring can support activity | Universal 5/10 rule | Clinical direct narrow |
| Sprague et al., 2021, DOI 10.1186/s40814-021-00792-5 | Patellar tendinopathy pilot | Pain-monitoring transfer is condition-specific | General recreational rule | Pilot |
| Tran et al., 2025, DOI 10.2519/jospt.2025.13253 | Pain-monitoring review | General evidence certainty is limited | Universal pain threshold | Review |
| Guo et al., 2017 | Calibration study of neural networks | Confidence must be calibrated to an outcome | Training decision confidence UI | Methods / adjacent |
| Hüllermeier & Waegeman, 2021 | Uncertainty taxonomy | Distinguishes uncertainty types | Product-specific safety efficacy | Methods review |
| Lee & See, 2004 | Human-automation trust review | Appropriate reliance matters | App confidence metric | Conceptual review |
| Kompa et al., 2021 | Abstention in medical ML | Uncertainty can trigger review/abstention | Consumer training outcome benefit | Adjacent |
| MacroFactor public algorithm documents | Product documentation | Smoothing, feedback loops, data-sufficiency precedent | Scientific validation of Hybrid Engine | Product precedent |
| RTS / Juggernaut public coaching material | Named coaching methodology | RPE/RIR, fatigue, wave, and feedback-loop precedents | Independent validation of complete systems | Coaching precedent |

## Appendix 2. Source metadata and current links

The principal sources are discoverable through the identifiers below. URLs are included for retrieval, not as a claim that a landing page is itself the evidence.

- Ratamess NA, Alvar BA, Evetoch TK, et al. Progression Models in Resistance Training for Healthy Adults. *Medicine & Science in Sports & Exercise*. 2009;41(3):687–708. DOI: 10.1249/MSS.0b013e3181915670. PubMed: https://pubmed.ncbi.nlm.nih.gov/19204579/
- American College of Sports Medicine. Resistance Training Prescription for Muscle Function, Hypertrophy, and Physical Performance in Healthy Adults: An Overview of Reviews. 2026 position stand. ACSM summary: https://acsm.org/resistance-training-guidelines-update-2026/
- Plotkin DL, et al. Progressive overload without progressing load? The effects of load or repetition progression on muscular adaptations. 2022. PubMed: https://pubmed.ncbi.nlm.nih.gov/36199287/. Full text: https://pmc.ncbi.nlm.nih.gov/articles/PMC9528903/
- Hostler D, et al. The effectiveness of 0.5-lb increments in progressive resistance exercise. 2001. PubMed: https://pubmed.ncbi.nlm.nih.gov/11708713/
- Buskard ANL, et al. Optimal Approach to Load Progressions during Strength Training in Older Adults. *Medicine & Science in Sports & Exercise*. 2019;51(11):2224–2233. DOI: 10.1249/MSS.0000000000002038. PubMed: https://pubmed.ncbi.nlm.nih.gov/31107348/
- Hickmott LM, et al. The Effect of Load and Volume Autoregulation on Muscular Strength and Hypertrophy: A Systematic Review. 2022. Full text: https://pmc.ncbi.nlm.nih.gov/articles/PMC8762534/
- Rogerson D, et al. Deloading Practices in Strength and Physique Sports. *Sports Medicine - Open*. 2024;10:26. DOI: 10.1186/s40798-024-00691-y. Full text: https://pmc.ncbi.nlm.nih.gov/articles/PMC10948666/
- Pancar Z, et al. Effects of deload periods in resistance training on muscle hypertrophy and strength endurance in untrained young men. *Scientific Reports*. 2026;16:10299. DOI: 10.1038/s41598-026-40612-5. Full text: https://www.nature.com/articles/s41598-026-40612-5
- Coleman M, et al. Gaining more from doing less? The effects of a one-week deload period. 2024. Full text: https://pmc.ncbi.nlm.nih.gov/articles/PMC10809978/
- Bell L, et al. Integrating Deloading into Strength and Physique Sports Training. 2023. Full text: https://pmc.ncbi.nlm.nih.gov/articles/PMC10511399/
- Addleman JS, et al. Heart Rate Variability Applications in Strength and Conditioning. 2024. Full text: https://pmc.ncbi.nlm.nih.gov/articles/PMC11204851/
- Grgic J, et al. Test-Retest Reliability of the One-Repetition Maximum Test: A Systematic Review. 2020. PubMed: https://pubmed.ncbi.nlm.nih.gov/32681399/
- Silbernagel KG, et al. Continued sports activity using a pain-monitoring model during rehabilitation in patients with Achilles tendinopathy. DOI: 10.1177/0363546506298279.

## Appendix 3. Inherited project evidence bundle and design record

The following appendices preserve the project’s prior evidence bundle, mechanism audit, and product design record. They remain part of the project audit trail. The reconciled conclusions in the main body govern where earlier wording was broader than the final evidence permits.

''')

def add_generated_material(base):
    # Preserve the existing evidence bundle and design decisions as a searchable audit trail.
    inherited = []
    for name in [
        "hybrid_adaptive_evidence_bundle_2026-08-01.md",
        "hybrid-engine-evidence-audit-mechanisms-7-8.md",
        "design.md",
    ]:
        path = ROOT / name
        if path.exists():
            inherited.append(f"\n\n### Preserved file: {name}\n\n" + path.read_text(encoding="utf-8"))
    return base + SUPPLEMENT + EXPANDED_MONOGRAPH + "".join(inherited)

SUPPLEMENT = dedent(r'''

## Appendix 4. Detailed decision-state catalogue

The following catalogue turns the research conclusion into observable product behaviour. It is intentionally more specific than the evidence itself. These states are implementation choices, not claims that a study has validated each one. Their value is that they prevent the product from collapsing every difficult day into one opaque readiness number.

### State: approved

Use `approved` when the planned prescription is compatible with the current safety state, the athlete has enough comparable evidence for the requested action, and no higher-priority conflict is present. Approval does not mean the athlete is guaranteed to complete the work. It means the prescription is reasonable to attempt. The app should still monitor technique, symptoms, and unexpected RPE during the session. An approved decision can become held, reduced, substituted, or blocked after new information appears.

### State: held

Use `held` when the current dose remains reasonable but the evidence does not justify increasing it. This is the normal result for noisy or incomplete data. The athlete should not interpret a hold as a failure. A hold preserves the planned exposure and protects the next decision. The explanation should name the missing or conflicting evidence, such as an unusually coarse plate jump, missing RPE, a single poor set, or a wearable signal that conflicts with normal performance.

### State: repeated

Use `repeated` when the athlete should repeat the current exposure to obtain a comparable observation. This differs from held: held may complete the current session without a progression change, while repeated explicitly schedules the same or nearly same exposure again. Repetition is appropriate after one non-painful miss, an unfamiliar movement, a large technique change, or an uncertain return from a short gap. It should not be used to pressure an athlete through pain.

### State: reduced

Use `reduced` when the planned purpose can be preserved with a lower stress dose. The reduction should identify its lever: load, sets, repetitions, density, complexity, or conditioning intensity. The engine should avoid changing multiple levers unless the state is systemic or safety-related. A reduced strength session can retain the primary lift and remove accessories. A reduced conditioning session can retain aerobic work while removing intervals. The user should be told what was protected.

### State: substituted

Use `substituted` when the original movement or modality is unsuitable but the session purpose remains achievable through an approved alternative. The substitution must retain the intended movement or energy-system role as closely as possible and must be recorded as a different exposure for trend purposes. A substitution is not equivalent data unless the exercise library declares it equivalent for a specific decision. The engine should never silently combine a substituted row with the original exercise’s strength trend.

### State: calibration

Use `calibration` after a meaningful training gap, a major programme change, a new exercise, a new device, or another event that makes the current estimate unreliable. Calibration is not punishment and is not a diagnosis of detraining. It is a deliberate information-gathering state. The first exposure should be conservative enough to reduce unnecessary failure risk and informative enough to observe technique, RPE, symptoms, and completion. Exit criteria must be based on successful comparable evidence, not on a hidden time counter alone.

### State: recovery

Use `recovery` when the session purpose is temporarily changed to restore capacity without abandoning the training habit. Recovery can mean easy Zone 2, mobility, technique, a short minimum viable session, or complete rest when appropriate. It is not automatically triggered by one low HRV value. It should be selected when direct performance, symptoms, recent training stress, and context support the conclusion that the planned dose is currently too expensive.

### State: blocked

Use `blocked` for a hard safety condition, a movement-specific pain block, invalid data that could make the action unsafe, or a user action that requires professional review. A block is not the same as a low-confidence recommendation. The product must make clear what is blocked, what safe alternative exists, and what information or professional input is needed before the blocked path can resume. The app must not allow a motivational override to bypass the state.

### State: needs_input

Use `needs_input` when one missing piece of information would materially change the decision and can reasonably be collected from the athlete or coach. Ask one useful question rather than presenting a questionnaire. Examples include whether an incomplete set was stopped by pain or fatigue, whether a substitution was intentional, or whether the load was entered in kilograms or pounds. Missing optional data should lower confidence without blocking the session.

### State: needs_clinical_review

Use `needs_clinical_review` for red-flag symptoms, persistent or worsening pain, neurological symptoms, chest pain, fainting, severe breathlessness outside the intended conditioning stimulus, or another configured reason to leave autonomous programming. The product should not diagnose. It should state that the training decision is outside the engine’s authority and advise appropriate professional or urgent care according to the configured safety copy.

## Appendix 5. Acceptance matrix for the adaptive layer

The matrix below is the minimum test inventory before the adaptive engine is considered complete. Each test should run against the pure rules engine and, where relevant, through the athlete UI, coach assignment boundary, offline persistence, sync merge, and rehydrated session state.

''')

def make_acceptance_matrix():
    groups = {
        "Load arithmetic": [
            "A 25 kg anchor with a 2.5% target and 2.5 kg equipment step must display an actual 10% jump and choose hold or fallback, never label it 2.5%.",
            "A 50 kg anchor with a 2.5 kg step must calculate an actual 5% jump and apply the movement cap before promoting.",
            "A 100 kg anchor with a 2.5 kg step must calculate an actual 2.5% jump without floating-point drift.",
            "A 180 kg anchor with a 2.5 kg step must calculate an actual 1.388...% jump and retain the unrounded calculation in the audit record.",
            "A zero or negative anchor must fail validation and return needs_input rather than producing an infinite or negative percentage.",
            "A missing equipment step must not be treated as zero; the engine must hold or request gym setup information.",
            "A unit conversion from pounds to kilograms must be explicit, versioned, and visible in the decision inputs.",
            "A candidate load below the anchor must not be selected as a promotion even if it is the closest available equipment value.",
            "A plate step that exceeds the movement cap must trigger fallback without deleting the success evidence.",
            "Two candidate loads equally close to target must be resolved by a deterministic tie-break rule.",
            "A machine with a fixed 5 kg step must not inherit a barbell microplate configuration without an explicit equipment identity.",
            "A changed bar weight must invalidate a direct comparison unless the exercise configuration declares the change safe.",
            "A candidate load must preserve the programmed unit and display unit conversion only as presentation when appropriate.",
            "The engine must store target percentage and actual percentage separately when rounding changes the physical jump.",
            "A progression percentage outside configured safety bounds must be rejected or require coach authorisation.",
        ],
        "Exposure validity": [
            "Changing from a flat bench to an incline bench must create a separate variation identity and not increment the flat-bench success streak.",
            "Changing the prescribed repetition range must mark the exposure non-comparable for the old progression rule.",
            "A set completed with a pain flag must not count as a normal successful exposure.",
            "A set completed with missing repetitions must not be inferred from the target value.",
            "An exercise substitution must preserve the session purpose but remain separate trend data unless equivalence is configured.",
            "A large rest-interval change must lower comparability when the exercise rule declares rest material.",
            "A deleted or edited historical set must create an auditable correction rather than silently changing the old decision.",
            "A duplicate offline submission must be idempotent and must not create a second success exposure.",
            "A session marked incomplete must not be promoted as complete solely because all entered sets appear valid.",
            "A coach target update must not overwrite athlete-owned actual load or repetitions.",
            "An athlete note saying ‘felt easy’ without an RPE value must remain qualitative context, not a numeric RPE.",
            "A set entered after the session date must retain its actual timestamp and not be backfilled as a normal exposure without an audit note.",
            "A changed exercise video or cue alone should not invalidate a comparable exposure if the movement identity remains stable.",
            "A new device source must be marked as a data-source change until its measurement behaviour is known.",
            "A coach-published target snapshot must remain reproducible even if the coach later edits the template.",
        ],
        "Promotion and holds": [
            "One valid success must hold the next load when the confirmation count is two.",
            "Two comparable successes must permit promotion only if no pain, calibration, or unresolved conflict state exists.",
            "A successful exposure with missing RPE may count toward completion but not toward high-confidence progression when RPE is required by the movement rule.",
            "A strong performance after unusually high conditioning stress must not automatically increase both strength and conditioning dose.",
            "A good readiness score with poor performance must produce hold or investigation, not automatic escalation.",
            "A low wearable recovery score with normal direct performance must not force a reduction by itself.",
            "A successful repetition fallback must be logged as progress without falsely claiming a load promotion.",
            "A progression decision must change only one primary lever unless a composite state explicitly allows more.",
            "A promotion must carry engine version, reason codes, source observations, and actual rounded jump.",
            "A later engine version must be able to replay the old decision under the old version without rewriting history.",
            "A coach override must be recorded as an override and not masquerade as automatic engine approval.",
            "An athlete override must not bypass a pain block or clinical-review state.",
            "A hold should not reset the success streak unless the exposure was invalid or the rule explicitly says it does.",
            "A failed repetition fallback must not create an additional hidden load reduction.",
            "A long period of holds must trigger review or data-quality guidance rather than endless silent repetition.",
        ],
        "Misses and reductions": [
            "One comparable non-painful miss must hold and record context rather than automatically reduce.",
            "Two comparable misses must be identifiable as the exact exposures that formed the heuristic evidence.",
            "A pain-related miss must not increment the normal performance-miss streak.",
            "A miss after an unusually hard interval session must carry conditioning context into the decision.",
            "A miss caused by equipment failure must be invalid for progression and not treated as athlete decline.",
            "A cross-session reduction must calculate from the last successful opening anchor, not from the walked-down effective load.",
            "A 5% reactive reduction must display the rounded load and the intended unrounded value.",
            "A systemic decline state must be able to reduce volume without changing the load anchor silently.",
            "A local movement miss must not automatically reduce unrelated exercises in the same session.",
            "A whole-session decline must not be inferred from one accessory miss.",
            "A successful calibration exposure must not be counted as a normal promotion unless the calibration rule permits it.",
            "Repeated misses separated by a long gap must not be combined as consecutive comparable misses.",
            "A missed session with no performed sets must be distinct from a failed session with attempted sets.",
            "The engine must never stack missed work into the next session without an explicit, safe programme rule.",
            "A user-requested ‘make-up’ session must preserve recovery spacing and return a reasoned alternative if unsafe.",
        ],
        "Pain and safety": [
            "Sharp pain during a movement must immediately leave normal fatigue logic and enter the configured symptom pathway.",
            "Chest pain, fainting, or severe unexplained dizziness must block autonomous training decisions.",
            "Delayed soreness must be stored separately from acute movement pain.",
            "A pain score alone must not be treated as a diagnosis or universal permission to continue.",
            "A pain block must apply to the affected movement or pattern without unnecessarily blocking unrelated safe work.",
            "A substitution after pain must be recorded as substitution data, not silently merged with the original trend.",
            "A pain block must persist across reload and offline restart until a valid update changes the state.",
            "An athlete pressing ‘continue anyway’ must not bypass a hard safety state.",
            "Clinical-review copy must avoid diagnosis, certainty, and false reassurance.",
            "The engine must preserve the user’s symptom description even if the coach later changes the planned exercise.",
        ],
        "HRV and context": [
            "A missing HRV value must be represented as missing, not as a normal score or zero.",
            "An HRV value outside the athlete’s personal baseline must lower confidence only when protocol and timestamp are valid.",
            "HRV alone must never increase load or volume.",
            "A low HRV plus normal performance must remain eligible for the planned purpose with progression held if configured.",
            "A low HRV plus poor performance and high recent conditioning load may enter conservative mode.",
            "A high HRV value must not authorise a risky progression when technique or pain is poor.",
            "Two wearable devices with incompatible values must create a conflict state rather than an averaged fiction.",
            "A stale recovery score must be marked stale and must not control a same-day decision.",
            "Sleep, stress, illness, and work context must remain distinct fields even if they contribute to a summary.",
            "The explanation must identify whether the recommendation was driven by direct performance or advisory wearable context.",
        ],
        "Layoff and calibration": [
            "A return after a configured meaningful gap must enter calibration rather than use an invented percentage equation.",
            "A one-day scheduling miss must not automatically trigger the same calibration state as a multi-week gap.",
            "Calibration must cap unnecessary failure exposure while retaining enough work to observe current ability.",
            "A calibration session with pain must remain in the safety pathway and not exit calibration as a normal success.",
            "A successful calibration exposure must record why it was not used as a normal anchor.",
            "Calibration exit must be deterministic and based on stated evidence, not an invisible timer.",
            "Different modalities must have separate layoff rules and must not share a single cross-modality reduction percentage.",
            "An illness return must not be treated as an ordinary training gap when symptoms are still active.",
            "An equipment change after a layoff must extend or restart calibration when comparability is lost.",
            "The user must see that calibration protects information quality and is not a punishment for time away.",
        ],
        "Data, sync, and ownership": [
            "Local writes must be durable before cloud sync is attempted.",
            "The same session submitted twice from offline replay must merge idempotently by stable identifiers.",
            "Coach target fields must never be overwritten by athlete result fields during reconciliation.",
            "Athlete result fields must never be erased because a coach publishes a later target snapshot.",
            "A deletion tombstone must prevent an old device from resurrecting a removed assignment.",
            "Derived estimates must retain the raw observations and engine version that produced them.",
            "A cloud conflict must be visible and resolved by the declared merge policy, not last-write-wins everywhere.",
            "An assignment must be self-contained enough for the phone to use it if the coach UI is unavailable.",
            "WHOOP secrets and provider credentials must never appear in browser storage or cached static assets.",
            "Authenticated function routes must not be cached by the service worker.",
            "A corrupted backup must fail closed and leave the current local state intact.",
            "Exported data must contain enough metadata to reproduce a decision without private provider tokens.",
            "Reset-local-data must be explicit, confirmed, and must not silently delete cloud history.",
            "The adaptive layer must remain usable offline with an honest degraded-data state.",
            "Coach and athlete apps must agree on the same exercise identity and prescription measurement schema.",
        ],
        "User experience and explanations": [
            "Every automatic change must have a plain-language explanation generated from structured reason codes.",
            "A hold message must say what was preserved and what evidence would permit progression.",
            "A reduction message must identify the changed lever and the reason for choosing that lever.",
            "A calibration message must explain that the engine is gathering current evidence rather than assuming detraining.",
            "A blocked message must distinguish safety limitation from missing optional data.",
            "The app must not show an unqualified confidence percentage without outcome, horizon, and data context.",
            "The user must be able to inspect the observations behind a decision without reading raw database JSON.",
            "The user must not be shamed for missing a session or asking for a shorter option.",
            "A minimum viable session must state which part of the original purpose was preserved.",
            "A coach must be able to see whether a target changed automatically, manually, or through athlete feedback.",
            "The app must avoid physiological explanations that exceed the evidence or imply diagnosis.",
            "The interface must make the next useful action obvious even when the answer is hold or rest.",
            "Reduced-motion and accessible touch targets must remain intact in the live logger.",
            "A user can reject an optional recommendation without losing the underlying record.",
            "History must show the original prescription and actual result separately.",
        ],
    }
    out = []
    n = 1
    for group, items in groups.items():
        out.append(f"### {group}\n")
        for item in items:
            out.append(f"{n}. {item}")
            n += 1
    return "\n".join(out)

SUPPLEMENT += make_acceptance_matrix()

SUPPLEMENT += dedent(r'''

## Appendix 6. Operational protocols for evidence-safe iteration

### Protocol A: changing a constant

No constant should be changed merely because a user had one unusually good or bad session. A proposed change must identify the current constant, its evidence status, the expected benefit, the possible failure mode, the affected movement classes, and the replay fixtures that will be rerun. The change should be represented as an engine-version update so old decisions remain reproducible. If the change affects a safety gate, it requires a separate review from ordinary progression tuning.

### Protocol B: adding a new data source

When adding a wearable, heart-rate bridge, machine sensor, or coach-entered field, first define its measurement semantics. What exactly is measured? At what time? With what missingness? How is it normalised? What population or device validation exists? Which decisions may use it, and which decisions may it never control? A new source begins in advisory mode. It may graduate to a stronger role only after the app has enough within-user data and a documented validation result. The presence of a number is not proof of usefulness.

### Protocol C: interpreting a surprising athlete result

When an athlete dramatically exceeds a target, do not immediately increase all future work. Check whether the target was entered correctly, whether the exercise and equipment were comparable, whether the repetitions were valid, whether the athlete misread RPE, and whether the result persists. When an athlete dramatically underperforms, check the same categories plus sleep, illness, conditioning, pain, and work stress. The first response should usually be classification, not storytelling. The engine can say “unexpected result; hold and verify” without inventing a mechanism.

### Protocol D: handling user overrides

Overrides are valuable evidence because they show where the engine’s default does not fit the athlete’s reality. They must not be hidden. Store the requested action, the engine action, the reason for the override, whether a safety state was bypassed, and the result. An override may change a non-safety prescription, such as removing an accessory because of time. It must not bypass a pain block, clinical-review state, or red-flag stop. Repeated overrides of the same type should become a product-review signal rather than a reason to silently adapt the rule.

### Protocol E: distinguishing adherence from efficacy

If a user completes more sessions under the adaptive engine, that is an adherence or usability signal. It is not by itself proof that the engine produces more strength or hypertrophy. If strength improves, the engine still cannot claim causation without a suitable comparison. Product validation should report adherence, completion, user comprehension, safety-state integrity, and performance trends separately. The system is allowed to be useful before it is proven superior, but it is not allowed to call usefulness superiority.

### Protocol F: closing a research question

A research question is closed for implementation when the evidence has been searched to a reasonable depth, direct evidence and gaps have been stated, the product decision is explicit, the decision’s provenance is recorded, and the remaining uncertainty is assigned to validation rather than endless browsing. The question can reopen if a new direct trial appears, a safety incident occurs, or product data shows the heuristic is systematically inappropriate. “Closed” means ready to build and test; it does not mean eternally beyond revision.

### Protocol G: audit log minimum

Every adaptive decision should retain the engine version, timestamp, user and exercise identity, session purpose, anchor load, opening load, effective load, target increment, actual increment, equipment step, success/miss evidence, RPE/RIR, symptoms, recent conditioning summary, wearable status, missing fields, conflicts, reason codes, chosen state, and next recheck point. The user interface need not show every field at once, but the record must exist for replay and safety review.

### Protocol H: release gate

The adaptive layer is not release-ready until static checks, browser smoke tests, offline reload tests, sync merge tests, security checks, and the acceptance matrix pass. A polished screen is not evidence that the engine works. A passing unit test is not evidence that the coach and athlete data boundary is intact. A green deployment is not evidence that a WHOOP secret is protected. Each layer has its own acceptance criteria.

### Protocol I: what to do when evidence conflicts

Conflicting studies should not be reduced to a vote. First identify whether they tested the same outcome, population, exercise, dose, and deload definition. Then determine whether the conflict is real or caused by different questions. If the conflict remains, choose the lower-risk product policy, expose the uncertainty, and preserve an experiment path. A conflict is not a failure of science; it is information about where the product should avoid false certainty.

### Protocol J: the human handoff

The engine should know when not to decide. Persistent pain, red flags, complex medical history, severe illness, or a request for rehabilitation guidance should leave the autonomous path. A coach or clinician can provide a target, but the data contract must record that the decision was human-authored or clinically informed. The system should never imply that a coach override became objective truth merely because it was entered through the same interface.

## Appendix 7. Final engineering checklist

Before the project is called complete, confirm the following in the repository and deployed application:

- The adaptive rules are pure, deterministic functions with versioned configuration.
- The current target, actual result, and derived estimate are separate data types.
- A flat kilogram rule is not used globally across exercises.
- Target and actual percentage changes are both calculated and displayed correctly.
- Equipment rounding cannot silently exceed the configured movement cap.
- Repetition fallback exists where a safe load jump cannot be expressed.
- One miss holds and gathers context; it does not trigger an unexplained large reduction.
- Repeated comparable deterioration is distinguished from pain, illness, and invalid data.
- Reactive reduction is calculated from the last successful opening anchor.
- Within-session correction is not silently compounded into cross-session reduction.
- HRV is advisory and cannot independently increase load or override pain.
- Layoff return enters calibration and has no universal time-off percentage equation.
- Pain is separate from fatigue and has a persistent block or modification path.
- Every decision has reason codes, evidence status, missing data, conflicts, and next recheck.
- Coach target fields and athlete result fields cannot overwrite one another.
- Local-first persistence survives reload, screen lock where supported, and offline use.
- Sync merges stable records without resurrecting deleted work.
- Secrets remain server-side and authenticated routes are not cached.
- Historical raw observations can replay derived decisions under the original engine version.
- The app communicates holds and reductions without shame or false certainty.
- The repository includes this dossier or an equivalent decision record next to the implementation.

The project is closed when these checks are either passed or deliberately recorded as a known release limitation with an owner and next action. A known limitation is acceptable. An invisible limitation is not.

## Appendix 8. Sensitivity analysis: what changing the heuristic would do

Because the literature does not identify an optimal increment, the product should understand the practical consequences of plausible settings. The following analysis is not a physiological prediction. It is a control-policy sensitivity analysis: how often would the engine create a large hardware-driven jump, how often would it hold, and how much historical confirmation would be needed before the athlete sees a different prescription?

At a 20 kg anchor, a 2.5 kg equipment step is a 12.5% jump. At a 25 kg anchor it is 10%. At 30 kg it is 8.33%. At 40 kg it is 6.25%. At 50 kg it is 5%. At 60 kg it is 4.17%. At 80 kg it is 3.125%. At 100 kg it is 2.5%. At 140 kg it is 1.79%. At 180 kg it is 1.39%. A single flat step therefore behaves like several different prescriptions. The engine’s cap is needed even if the target percentage is changed later.

If the target is reduced from 2.5% to 2%, more heavy-lift promotions will be expressed by the same available step because the step is already above the target, while more light movements will fall into the fallback state because the next physical jump remains too large. If the target is increased to 5%, more medium and heavy movements will wait for a larger jump or use repetition progression. That may be appropriate for a strength-focused block, but it will also delay physical load changes on machines with coarse steps. If the target is increased to 10%, most light movements will exceed the small-movement cap and most heavy movements will be vulnerable to a larger-than-necessary stress change. The engine should therefore treat target percentage and actual jump cap as separate controls.

The same sensitivity applies to the reactive reduction. A 5% reduction from 100 kg is 95 kg; a 10% reduction is 90 kg. On a 25 kg movement, the same settings produce 23.75 kg and 22.5 kg before rounding, which may be impossible with standard equipment. Rounding 23.75 kg to 22.5 kg creates a 10% physical reduction, while rounding to 25 kg creates no reduction. The app must display the physical load that will be used and the intended reduction so that a coach can see when hardware has distorted the nominal policy.

The safest first release is therefore not the one with the most precise-looking percentages. It is the one that makes distortion visible. For every load change, the audit record should answer: what was the anchor, what target change was requested, what hardware steps were available, what actual change was applied, whether the change exceeded the cap, and what fallback was used. This lets later validation evaluate the policy in the form the athlete actually experienced rather than in the abstract percentage the software intended.

## Appendix 9. What a decisive future study would need to test

The central unanswered question is experimentally tractable. A decisive study would recruit participants with sufficient training experience to make progression meaningful, stratify or randomise by sex and training status, and assign comparable exercises to progression policies that differ in increment size. The trial would need at least three arms, such as 2.5%, 5%, and 10%, with the trigger and target repetition range held constant. It would need to report actual load changes, not only intended rules, because plate granularity can make the arms converge or diverge.

The study should separate upper-body and lower-body movement classes and record the relative size of each physical jump. It should use familiar exercises, standardise range of motion and rest, record RPE/RIR and technique, and report adherence, missed sets, repeated exposure count, time to stall, strength, hypertrophy, and adverse events. It should pre-register how a successful exposure is defined and whether one, two, or three confirmations are required. Otherwise the study may answer the wrong question while appearing to answer the right one.

For deloading, the decisive design would compare a hold-only policy, a load-reduction policy, a volume-reduction policy, and a combined policy after a clearly defined and corroborated deterioration state. It should distinguish reactive from planned deloads, use trained participants, include a hybrid or concurrent-training condition, and report subsequent performance, adherence, perceived recovery, pain events, and training quality. A study that simply inserts a one-week cessation period in untrained people cannot settle the product’s reactive rule.

For HRV, the decisive design would compare a fixed programme, an HRV-only programme, and a multi-signal programme with pre-specified thresholds and missing-data handling. It would need to standardise measurement protocol, use within-user baselines, and test whether the policy improves meaningful outcomes rather than merely changing the number of hard sessions. The strength outcome should be separated from recovery or adherence outcomes.

Until such studies exist, the engine should not wait for certainty that the field does not currently possess. It should implement the conservative controller, label the constants, preserve the raw data, and make future comparison possible.

## Appendix 10. Why this dossier is sufficient to close the project

The purpose of the research phase was not to find a sentence that could be pasted into code as an unquestionable answer. It was to determine whether the proposed rules were supported, unsupported, or supportable as explicit engineering choices. That work is complete. The report has identified the strongest direct sources, traced the ACSM recommendation to its prescription lineage, corrected the overextension of Plotkin, located the closest microloading evidence, reviewed the available deload evidence, and separated the strength-specific questions from the hybrid-conditioning and product-safety questions.

It has also answered the uncomfortable questions that a less careful report would hide. There is no known direct percentage-band trial that proves 2.5% is better than 5% or 10%. There is no validated universal consecutive-failure threshold. There is no validated time-off-to-load equation. HRV is not a complete readiness gate. A pain-monitoring threshold from a specific rehabilitation protocol cannot become a universal consumer rule. A confidence badge is not calibration. These are not gaps that can be solved by writing more assertive prose.

The project can now move from research to implementation because each gap has a bounded response. The absence of a direct increment trial becomes a configurable percentage target, equipment cap, repetition fallback, and validation log. The absence of a deload threshold becomes repeated corroborated deterioration, a visible heuristic count, separate local and systemic pathways, and no hidden compounding. The absence of a layoff equation becomes calibration. The limits of HRV become advisory use. The limits of pain evidence become symptom separation and a safety handoff. The limits of confidence scores become action-linked evidence states.

This is the correct meaning of “evidence-informed” for a software system. It does not mean every branch was discovered by a randomised trial. It means the system knows which branches were studied, which are transferred cautiously, which are inherited from coaching practice, and which are product decisions made under uncertainty. It preserves enough raw data to be corrected later. It makes its own behaviour testable. It tells the athlete what it knows and what it does not know.

The remaining work is ordinary product work: encode the state machine, connect it to the existing local-first data boundary, write tests, run the app in the browser, inspect mobile and coach flows, and measure real behaviour. Reopening the same literature search before doing that would not increase the quality of the first release. The research phase has supplied the operating boundaries. The build phase must now prove that the implementation respects them.

Final decision: ship the bounded deterministic adaptive layer with the constants and labels in this dossier, then let controlled product validation—not endless argument—determine whether a later version should move those constants.

One final practical point deserves emphasis. A training engine does not become more intelligent by making a larger number of adjustments. It becomes more intelligent by making fewer unjustified adjustments and by learning from the ones it does make. A system that progresses every time the athlete reports a good set may appear responsive while quietly creating unstable load jumps. A system that deloads after every imperfect day may appear cautious while quietly destroying momentum. The controller must therefore be judged by proportionality: did it respond to the size and quality of the evidence, did it preserve the intended training purpose, and did it leave the next decision clearer than the last one?

That standard also protects the athlete from the psychological damage of over-interpretation. A hold is not a verdict on capability. A calibration session is not a declaration of weakness. A reduction is not punishment. A blocked movement is not the end of training. By making these states explicit, the app can support adherence without lying about certainty. This is especially important for a hybrid athlete whose capacity changes with work, sleep, conditioning, stress, and the ordinary unpredictability of life.

The final dossier is therefore both a research conclusion and a product boundary. It says what the engine can responsibly do now, what it must not claim, and what evidence would be required to change the rule later. That is enough to close the research phase and begin implementation with a clean conscience.

In short, the engine now has a defensible centre of gravity: direct performance outranks a single wearable signal; repeated comparable evidence outranks one dramatic session; safety and symptoms outrank training ambition; raw observations outrank derived scores; and explicit uncertainty outranks confident-looking folklore. Those priorities should remain stable even if the visible interface, programme templates, or future data sources change.

That stability is the final safeguard against scope drift. New features may add information, but they should not quietly promote an advisory signal into authority or convert a temporary heuristic into a permanent scientific claim. Any future change should pass through the same contract: identify the evidence, name the uncertainty, define the action, preserve the audit trail, and test the edge cases before release.

The next action is therefore implementation, not another debate about the decimal place. Start with the smallest coherent controller, keep every decision reversible, and use the athlete’s actual logged experience to decide whether the defaults deserve refinement.

That is the practical finish line: an honest rule, a clear boundary, a reproducible record, and a product that can improve without rewriting what happened.

No further research pass is required for the initial build. Future evidence belongs in versioned change records, not in an endless loop that prevents the system from being tested.

This document is the final research baseline for the first release and the reference point for all later rule changes.
''')

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_font(run, name="Calibri", size=11, color=None, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_md_table(doc, lines):
    rows = []
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(set(c.replace("-", "").replace(":", "").strip()) == set() for c in cells):
            continue
        if all(c.replace("-", "").replace(":", "").strip() == "" for c in cells):
            continue
        rows.append(cells)
    if len(rows) < 2:
        return
    header, body = rows[0], rows[2:] if len(rows) > 2 else []
    def clean_inline(text):
        text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
        text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
        text = text.replace("`", "")
        return text
    header = [clean_inline(x) for x in header]
    body = [[clean_inline(x) for x in row] for row in body]
    # Long evidence tables are more readable as labelled records than as
    # dense four-column grids. This also prevents citation-heavy rows from
    # splitting across pages without their column headings.
    if len(header) >= 4 and max((len(cell) for row in body for cell in row), default=0) > 90:
        for idx, row in enumerate(body, 1):
            p = doc.add_paragraph()
            r = p.add_run(f"Evidence record {idx}")
            set_font(r, size=10, color="1F4D78", bold=True)
            for label, value in zip(header, row):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                r = p.add_run(f"{label}: ")
                set_font(r, size=9.5, bold=True, color="0B2545")
                r = p.add_run(value)
                set_font(r, size=9.5)
            doc.add_paragraph()
        return
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(6.5 / len(header))] * len(header)
    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_font(r, size=9, bold=True, color="0B2545")
    for row in body:
        if len(row) != len(header):
            continue
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].width = widths[i]
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            r = p.add_run(text)
            set_font(r, size=8.5)
    doc.add_paragraph()

def build_docx(md):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    # Quiet running furniture for a long formal research dossier.
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("THE Hybrid Engine  |  Final evidence dossier")
    set_font(r, size=8, color="667085")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Final close-out  •  16 August 2026")
    set_font(r, size=8, color="667085")

    lines = md.splitlines()
    i = 0
    first_title = True
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("| "):
            tbl = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl.append(lines[i])
                i += 1
            add_md_table(doc, tbl)
            continue
        if line.startswith("```"):
            code = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
            for idx, code_line in enumerate(code):
                r = p.add_run(code_line + ("\n" if idx < len(code)-1 else ""))
                set_font(r, name="Courier New", size=8.5, color="1F2937")
            i += 0
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(line[2:].strip())
            set_font(r, size=24 if first_title else 18, color="0B2545", bold=True)
            p.paragraph_format.space_after = Pt(10)
            first_title = False
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(line[2:].strip())
            set_font(r, size=10.5)
            i += 1
            continue
        if line[:3].isdigit() and line[3:5] == ". ":
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(line[5:].strip())
            set_font(r, size=10.5)
            i += 1
            continue
        p = doc.add_paragraph()
        # Basic inline emphasis without risking malformed Word runs.
        text = line.strip()
        parts = text.split("**")
        for idx, part in enumerate(parts):
            if not part:
                continue
            r = p.add_run(part)
            set_font(r, size=10.5, bold=(idx % 2 == 1))
        i += 1

    # Core properties and a simple first-page title treatment.
    doc.core_properties.title = "THE Hybrid Engine — Final Evidence Dossier and Implementation Lock"
    doc.core_properties.subject = "Evidence audit and deterministic implementation specification"
    doc.core_properties.author = "THE Hybrid Engine"
    doc.core_properties.comments = "Final project close-out document."
    doc.save(OUT_DOCX)

REPO_AUDIT = dedent(r'''

## Part VII. Repository-grounded implementation audit

### Audit scope and source boundary

This part records what the current repository contributes to the close-out decision. It is not a reconstruction from the product brief and it is not a claim that every planned surface is live. The inspected repository is `reflectprotect123-max/THE-HYBRID-ENGINE1`, public on GitHub, default branch `main`, inspected at the 16 August 2026 close-out point. The current tree is identified in the audit notes by commit `e07eeae8dc3d863a71272985a66b5b52b278bf8d`. The direct repository source is [THE-HYBRID-ENGINE1 on GitHub](https://github.com/reflectprotect123-max/THE-HYBRID-ENGINE1). File-level claims below should be rechecked against a later commit before release.

The repository changes the character of the project. Earlier research could only specify what a safe engine ought to do. This codebase shows which decisions have already been made, where they have been made, and which boundaries have been recognised by the author. It also shows where a good local rule can remain ineffective because the caller, product shell, or cross-package contract does not consume it. For that reason the audit separates five states: implemented and exercised; implemented but not proven effective in the product; designed but not wired; deliberately excluded; and unresolved.

### Architecture readout

The monorepo is organised around a shared engine rather than a single screen. `packages/engine` contains the legacy-compatible training model and a local merge boundary. The package has explicit modules for sessions, logging, lifts, autoregulation, conditioning, heart-rate interpretation, balance, insights, database sanitisation, cloud mapping, coach-to-athlete emission, numeric helpers, plate math, and storage. That decomposition is valuable because it puts arithmetic, persistence, and presentation in separable places. It also creates a verification obligation: every decision must have one authoritative calculation and every consumer must call it.

`packages/shared-core` holds versioned facts and cross-application contracts. `packages/whole-athlete-state` holds recovery, life context, constraints, and data quality. `packages/session-authoring` represents authoring boundaries. Nutrition is separated into core, engine, and adapter packages. The Android athlete application is the whole-athlete surface; the web application is the coach workspace. The README records that the earlier athlete web surface and several coordinator/auto-coach/strength-engine/conditioning-engine packages were deleted. This is a healthy reduction in duplicate orchestration, but it means old rows, migrations, and compatibility names must not be mistaken for active producers.

### What the engine already does well

The lift path has unusually clear provenance. `lift.ts` distinguishes the last completed non-warm-up working set from warm-ups, requires real repetitions, uses the same plan-anchored fold to compute the next opener, and banks the result as a replacement map rather than mutating settings in place. It protects against a late or restored session overwriting a newer record. `nextWorkingWeight` applies today’s recovery adjustment at read time, so a bad morning can ease the offered weight without erasing what the athlete previously earned. `openingLoadFor` makes the precedence rule explicit: today’s logged fact, authored percentage, earned weight, then honest blank/zero.

The conditioning path is equally important because it demonstrates restraint. `conditioning.ts` distinguishes formats and modalities, treats recovery as a prescription-time modifier rather than a reason to erase earned progression, excludes simulated sessions, refuses to earn progress when the heart-rate record contains no usable zone time, uses work time rather than rest-inclusive duration for the denominator, and keeps the HRR gate provisional rather than pretending it is validated. The one-session success path is not allowed to become a universal dose equation.

The adaptive strength layer in `packages/engine/src/adaptive/strength.ts` adds a read-only, explanation-oriented decision function rather than a second storage system. It requires three exposures before making a suggestion, uses the same exposure selection rules as lift progression, supports a repetition route before a load route, and returns an explicit `pause_insufficient_data` state when history is thin. Its result carries action, confidence, reason codes, safety state, data limitations, and an optional prescription. That contract is the right shape for a coach or athlete surface because the number is not detached from why it was offered.

`plates.ts` is a small module with a large safety implication. It sanitises inputs, searches a deterministic greedy-plus-correction space, and returns both the achievable load and the delta from the requested load. It refuses to pretend that an unachievable prescription was hit. Equipment-aware rounding is not cosmetic: a 2.5% target can become a 10% physical jump on a light lift, and a safe controller must know the difference.

### Whole-athlete state: model present, integration still decisive

The whole-athlete package has a clear schema. It represents readiness as a band with data quality, capacity separately for strength and conditioning, recovery debt, illness status, advisory HRV, constraints, and rationale. Pain and illness are hard constraints; low readiness, high recovery debt, high physical load, time limitation, and low energy availability are softer training constraints. HRV is explicitly advisory and does not by itself alter readiness or create a pain gate in the tests. Nutrition is passed as observed intake and estimated expenditure, not as a target or instruction.

The crucial close-out finding is not that this model is absent. It is that the repository documentation says the package raises pain and illness flags but, at the inspected point, nothing consumes them downstream. That is a wiring gap, not a research gap. A hard constraint that is never read by the session selector, prescription layer, logger, or coach review surface is only a typed intention. The project cannot be considered closed until the consumer path is proven with an integration test: create the flag, ask for a session decision, verify that normal progression is blocked or redirected, and verify that the explanation names the hard constraint.

### Verification and deployment boundary

The root scripts show a serious verification posture: typecheck, unit tests, ecosystem contracts, reachability, CSS-state checks, site build, CSP checks, React smoke, deploy smoke, and parity harnesses. The README also calls out offline reload, live sync behind an explicit environment gate, server-side WHOOP and Concept2 functions, encrypted provider tokens, and service-worker update behaviour. These are not training-science findings, but they are part of a safe adaptive engine because a correct rule that is stale, duplicated, unsynced, or bypassed by a browser cache is not correct in use.

The repository’s `.claude/agents` directory contains a read-only locator role and a review-only role, while the dispatching skill asks for independent domain passes followed by integration and a full suite. This dossier follows that discipline as a multi-pass audit of the code and evidence. No external agent output is being presented as if it ran; the source of the implementation findings is the repository itself. The distinction matters because orchestration metadata is not test evidence.

### Repo-grounded close-out judgement

The project has crossed the line from idea to implementable system. The core engine contains real guards, explicit precedence, persistence semantics, explanations, and tests. The evidence-supported product shape is therefore strong enough to freeze as a deterministic baseline. It has not crossed the line to “finished” merely because the code is extensive. The remaining closure work is integration proof: consume whole-athlete hard constraints, prove the adaptive decision output reaches the actual athlete surface, verify coach/athlete ownership through sync, and run the full verification command at the commit being released. Those are finite release gates. They are not invitations to reopen the entire research question.

### Repository evidence map

| Area | Current repository evidence | Close-out interpretation |
|---|---|---|
| Strength progression | `packages/engine/src/lift.ts`, `adaptive/strength.ts`, tests | Strong deterministic foundation; validate caller reachability |
| Within-session RPE | `autoreg.ts`, `fold.ts`, session tests | Use as a control signal, not a diagnosis |
| Conditioning | `conditioning.ts`, conditioning tests | Good separation of earn, prescribe, exclude, and hold |
| HRV/recovery | `hr.ts`, whole-athlete state, state tests | Advisory unless corroborated; missingness must remain visible |
| Pain/illness | whole-athlete state types/state | Hard flags exist; consumer integration is the major gap |
| Equipment | `plates.ts`, numeric helpers | Physical jump must be computed, displayed, and audited |
| Data boundary | `db.ts`, `cloud.ts`, merge/restore tests | Preserve raw facts and deterministic reconciliation |
| Product surfaces | Android athlete app plus web coach workspace | Verify one decision path reaches the athlete |
| Release proof | root `verify` and smoke/parity scripts | Run at release commit; do not infer from unit tests |

''')


# The following records are deliberately written as implementation questions, not
# invented literature findings. The renderer expands them into a long-form audit
# so the 100,000-word version remains searchable by mechanism, file boundary, and
# acceptance criterion rather than becoming a block of unstructured prose.
MONOGRAPH_PACKS = [
    ("Chapter 1 — Repository anatomy and ownership", [
        ("One engine, many surfaces", "The same training fact must have one authoritative calculation even when it appears in mobile logging, coach review, summaries, and sync payloads.", "The repository separates engine, shared-core, session-authoring, nutrition, mobile, and web packages.", "Duplicated formulas create silent disagreement between what the athlete sees and what is stored.", "Trace one completed lift from authored target through logger, local persistence, sync, recap, and coach view.", "Keep calculation in the engine and make every surface consume the typed output."),
        ("The local-first boundary", "Offline use is a product requirement, not a convenience, because the athlete may be training without a network connection.", "The README describes durable local storage, restore, merge, and live-sync checks.", "A cloud-first design can lose a completed session or make a stale target appear current.", "Kill the network after a completed set, reload, resume, sync later, and verify idempotent reconciliation.", "Local completion must be durable before remote acknowledgement is required."),
        ("Coach ownership versus athlete facts", "A coach-authored prescription and an athlete-entered result are different authorities.", "The repo names coach workspace and athlete app boundaries and contains cloud mapping and merge modules.", "Last-write-wins can let a later plan overwrite the actual load performed or erase a user correction.", "Submit conflicting target and result edits from separate clients and inspect the merged record and audit trail.", "Declare field ownership explicitly; never resolve semantic conflicts with timestamp alone."),
        ("Deleted orchestration packages", "Removed packages and historical database names must not be treated as live behaviour.", "The README records deletion of coordinator, auto-coach, strength-engine, conditioning-engine, and AI-prescription packages.", "A migration or compatibility writer can leave the impression that an old producer still exists.", "Search imports, build outputs, row writers, and route reachability for each deleted namespace.", "Retain compatibility reads only when they are documented and tested; remove ambiguous live paths."),
        ("Session authoring as a single surface", "There should be one place where a coach builds the canonical day/session representation.", "The README identifies `DayBuilder.tsx` as the authoring surface and rejects a second builder.", "Two builders drift in defaults, exercise identity, or progression metadata.", "Create the same session through every exposed route and compare the serialised assignment.", "One canonical authoring boundary reduces ambiguity and makes replay possible."),
        ("Package contracts as safety boundaries", "A package boundary is useful only if its types prevent forbidden data from crossing it.", "Nutrition context intentionally omits targets, while whole-athlete state exposes observations and constraints.", "A broad `any` or convenience field can reintroduce a forbidden instruction through a side door.", "Compile contract fixtures that attempt to pass a nutrition target into training state and expect failure.", "Use types, tests, and review rules together; types alone are not sufficient evidence."),
        ("Versioned namespaces", "A decision must remain reproducible after schemas and algorithms evolve.", "Shared-core and whole-athlete state expose schema/version concepts.", "Without versioning, a later reader can reinterpret an old field under new semantics.", "Persist an old snapshot, update the package, restore it, and compare the replayed decision under the recorded version.", "Version facts and rules separately when either can change the meaning of a decision."),
        ("Mobile as the athlete truth surface", "The screen under the bar is where an adaptive rule becomes behaviour.", "The repository has an Android athlete app while the athlete web surface was removed.", "A correct helper that is not called by the mobile logger is functionally absent.", "Start from `openDraft`, follow the weight/reps prefill, and prove the value comes from the canonical engine function.", "Prioritise reachability and parity at the actual athlete interaction point."),
        ("Web as coach oversight", "Coach review needs more context than a single recommendation number.", "The web app is described as a coach workspace with readiness, strength, conditioning, nutrition, and progression routes.", "A coach view that hides reason codes can turn a heuristic into an unchallengeable command.", "Verify that each changed prescription displays source, reason, confidence, and the underlying exposure.", "The coach surface should expose the decision record, not merely the final load."),
        ("Verification as part of the model", "Build and smoke checks protect the semantics of an adaptive system at runtime.", "The root `verify` script includes typecheck, test, reachability, build, CSP, React smoke, and deploy smoke.", "Passing unit tests does not prove that a route is reachable or that a service worker serves current code.", "Run the complete verification command in a clean checkout and archive its output with the release commit.", "Treat verification results as release evidence with a timestamp and commit identity."),
    ]),
    ("Chapter 2 — Strength progression and loading", [
        ("Opening load as a fact hierarchy", "The engine must resolve what to put on the bar without letting several valid sources disagree silently.", "`openingLoadFor` documents fold, authored percentage, earned weight, then honest none/bodyweight handling.", "A stale banked load can override a coach’s explicit percentage, while an untouched fold can hide all history with zero.", "Exercise a fresh movement, a percentage target, a banked target, a completed set, a bodyweight movement, and a missing-data case.", "One function returns kilogram value, message, and source together."),
        ("Last working set versus warm-up", "Warm-up data is useful for readiness but must not become the working-load anchor.", "`lift.ts` filters warm-up blocks and warm-up sets before banking movement progress.", "An empty-bar warm-up at easy RPE can teach the system an absurdly low working weight.", "Log warm-ups before a heavy movement and assert that the banked opener follows the working set only.", "Keep warm-up exclusion in the shared selection helper, not in individual screens."),
        ("Ramp exercises", "A ramped movement must compare the next opener with the same opener that the decision describes.", "The lift tests include ramped exercises and protect `from`, `to`, delta, and reps as a single semantic pair.", "Reading the last set for `from` makes a valid 100/110/120 ramp look like an unexplained 120-to-100 change.", "Use flat and ramped fixtures with easy, on-target, hard, and incomplete final sets.", "Every displayed delta must compare like with like; otherwise the explanation is numerically false."),
        ("Authored percentage loads", "A percentage of e1RM is a prescription, not a replacement for today’s observed result.", "`prescribedKg` resolves an authored percentage against the same e1RM used elsewhere and rounds to the configured increment.", "Parsing `80%` as eighty repetitions or allowing a percentage with no usable history can produce unsafe nonsense.", "Test `5 @80%`, `@80%`, `80%`, out-of-range percentages, no history, and a conflicting earned load.", "Keep the percentage syntax explicit and refuse to guess when the reference e1RM is absent."),
        ("Microloading and actual percentage", "A target increment must be evaluated after equipment rounding, not before.", "The repo has plate math and numeric helpers that expose achievable load and delta.", "A nominal 2.5% instruction can become a 10% physical jump on a light bar or an impossible plate combination.", "Compare nominal target, rounded candidate, actual percentage, and equipment delta across light, medium, and heavy anchors.", "Store both intended and achieved changes so the controller can choose hold or fallback honestly."),
        ("The repetition fallback", "When equipment cannot express a safe load step, progress can move through repetitions or another declared lever.", "The adaptive strength layer includes a rep route before a load route and checks the current displayed target.", "A rep fallback that is lower than the plan’s existing target becomes a disguised downgrade.", "Test low-load exercises, top-of-range repetitions, bodyweight movements, and a plan already asking for the proposed rep count.", "The fallback must be explicit in the decision record and must not masquerade as load progression."),
        ("One lever at a time", "Load, repetitions, sets, rest, and density are not interchangeable interventions.", "The strength layer separates `progress_load`, `progress_reps`, hold, and reduction actions.", "Changing two or three levers at once destroys attribution and can overspend recovery.", "Create a two-session success streak and verify exactly one primary prescription field changes.", "Composite changes require a separately named state and a stronger reason than ordinary success."),
        ("Stable anchors", "Progression should use a successful opening anchor rather than a failed, walked-down effective load.", "The lift comments distinguish what was earned across the session from what the next opener is reported against.", "Compounding a within-session back-off with a cross-session reduction can spiral downward.", "Run a hard final set that walks the fold down, then inspect the next session’s anchor and reduction calculation.", "Record anchor, effective load, and next offered load as separate fields."),
        ("Comparability of exercise identity", "A progression streak is meaningful only when the exposures represent the same movement and context.", "The engine normalises names and the adaptive layer selects a first comparable occurrence per session.", "An incline bench or substituted machine can falsely continue a flat-bench streak.", "Change variation, equipment, grip, range, and target while preserving the display name and assert that comparability breaks.", "Exercise identity must include the dimensions that matter for the rule, not only a label."),
        ("When to hold", "A hold is an active control output that preserves information and avoids overreacting.", "The adaptive contract includes a hold action with reason codes such as mixed results and already-earned load.", "Users may interpret hold as failure if the interface does not explain what was preserved.", "Trigger one success, one miss, conflicting wearable data, and insufficient history; verify distinct hold explanations.", "A good hold states the next evidence that would permit a change."),
        ("Three exposures before suggestion", "Sparse history should lower confidence and prevent the system from writing a confident prescription.", "`adaptive/strength.ts` defines a minimum exposure count and returns `pause_insufficient_data` below it.", "A single easy set can be a measurement error, novelty effect, or unusual day.", "Provide zero, one, two, and three valid exposures, including invalid and non-comparable sessions.", "Use a minimum data rule as a safety feature, not as a claim that three is physiologically optimal."),
    ]),
    ("Chapter 3 — RPE, RIR, and within-session control", [
        ("RPE as an observed control signal", "RPE can guide a next-set adjustment without becoming a diagnosis of fatigue or readiness.", "`autoreg.ts` parses target and felt effort while the fold applies a bounded plan-anchored adjustment.", "Treating one rating as a precise physiological measurement can create false confidence.", "Use missing, nonnumeric, boundary, and contradictory RPE values and verify safe fallback behaviour.", "Use RPE to control the local dose and preserve the uncertainty around what caused the rating."),
        ("Target RPE versus felt RPE", "The controller must compare what was felt with what was asked, not compare the athlete with the target itself.", "The lift implementation explicitly distinguishes `felt` from prescribed `rpe`.", "Judging every set against its own target makes every set look perfect and prevents adaptation.", "Run identical target sets with felt RPE below, at, and above target and inspect the verdict and next load.", "Store both values and name the deviation used by the rule."),
        ("RIR conversion", "RIR is useful when the interface or coaching practice uses it, but conversion must be explicit.", "The repo’s RPE/RIR utilities are tested as parsing and classification helpers.", "An unlabelled seven can mean RPE 7 or seven repetitions in reserve depending on the field.", "Test notation, units, missingness, and impossible values across the logger and persisted session.", "Never infer the scale from a bare number when the input contract can require a label."),
        ("The hard-set guard", "A set that misses the repetition floor is qualitatively different from a set that simply feels a little hard.", "The strength layer treats a missed rep floor as a miss and the fold can lock a hard exercise walk.", "A later easy back-off can incorrectly raise the movement again if the hard set does not lock the sequence.", "Create a missed target followed by an easy set and assert that the earned movement does not rise.", "A failed set must remain visible as the event that constrained the session."),
        ("Dead bands and noise", "Small deviations around target should not trigger constant up-and-down movement.", "The repo documents a plan-anchored walk and uses bounded adjustments rather than unlimited set-to-set movement.", "Without a dead band, normal rating noise becomes oscillation.", "Replay a sequence of target-minus-one, target, target-plus-one, and half-point ratings.", "Define the dead band as a product heuristic, record it, and tune it from replay data rather than intuition."),
        ("RPE with technical breakdown", "A numeric effort score cannot replace a technique or symptom flag.", "Whole-athlete state introduces hard pain constraints separate from ordinary readiness signals.", "A user may rate a painful set as easy while still making it unsafe to progress.", "Pair each RPE band with pain, form, and completion flags and assert that hard flags outrank effort.", "The safety state must be resolved before the ordinary effort controller is allowed to act."),
        ("Missing effort ratings", "Missing RPE should reduce what the engine claims to know, not silently become a normal rating.", "Adaptive explanations carry data limitations and distinguish low confidence.", "Defaulting missing effort to target creates artificial success streaks.", "Omit RPE from one set, one session, and all sessions; inspect bank, streak, and confidence outputs.", "A missing field may preserve the session but should block high-confidence escalation where effort is required."),
        ("RPE across exercise classes", "An RPE rule that works for a barbell lift may not transfer to a timed conditioning interval or isolation movement.", "The repository separates lift and conditioning modules and does not force one universal prescription path.", "Cross-domain reuse can turn modality-specific observations into a false common scale.", "Run matched ratings through strength, intervals, steady conditioning, and bodyweight paths.", "Keep shared parsing small and keep decision semantics modality-specific."),
        ("Athlete override of the number", "The athlete standing under the bar has information the app may not have.", "`nextWorkingWeight` keeps the field typeable and presents an eased offer rather than an immutable block.", "An override can become invisible data loss or a bypass of a hard safety state.", "Record a voluntary load change, a pain-related attempt to continue, and a blocked override separately.", "Allow non-safety override, preserve the reason, and never let motivation bypass a hard stop."),
        ("Explaining effort decisions", "The explanation must be generated from the same inputs that produced the number.", "`adaptive/explain.ts` reshapes existing outputs into typed reason codes without recomputing them.", "A second UI formula can display a reason that does not match the actual load.", "Compare every explanation to the underlying `WorkingWeight`, `Prescription`, or `AdaptResult` fixture.", "One computation, one structured result, one human-readable note."),
    ]),
    ("Chapter 4 — Conditioning and hybrid interference", [
        ("Format-specific conditioning", "Steady work, intervals, rowing, running, and mixed formats have different meaningful exposure variables.", "`conditioning.ts` keys progression by format and modality and avoids collapsing all conditioning into one level.", "A global conditioning level can reward an easy format and raise an unrelated high-intensity prescription.", "Progress each format independently, substitute modality, and inspect the progression map.", "Use format identity as part of the evidence key and expose it in the explanation."),
        ("Work time versus rest-inclusive time", "The denominator for conditioning quality should match the physiological question being asked.", "The repo comments explicitly use work-time denominator rather than total duration including rests.", "Including rest can make a hard interval session look artificially easy or a steady session look incomplete.", "Construct equal work with different rest and equal total duration with different work; compare completion classification.", "Keep work, rest, total duration, and zone time as separate raw facts."),
        ("Heart-rate data quality", "No usable heart-rate zone time should not become a quiet success or failure.", "`conAdapt` excludes sessions with no zone time from earn/deload logic.", "A missing sensor can be misread as low intensity and wrongly earn progression.", "Test no sensor, partial sensor, dropout, simulated record, and complete zone data.", "Missing cardiovascular data means no cardiovascular verdict, not zero effort."),
        ("RPE primary in short intervals", "Short intervals can make HR lag or obscure the intended work signal.", "The conditioning comments use RPE as primary for short intervals and HR as secondary/diagnostic.", "A universal HR gate can penalise a valid interval session or reward a delayed heart-rate response.", "Compare short intervals with identical work but different HR response and use direct effort and completion as the primary evidence.", "The evidence hierarchy must be format-specific and visible."),
        ("No HRR gate by default", "Heart-rate recovery is attractive as a marker but not automatically a validated training gate.", "The current code marks the HRR condition provisional and leaves it non-gating.", "A named metric can acquire authority merely because it appears precise.", "Replay records with high, low, missing, and delayed HRR and verify that the current rule does not invent a hard stop.", "Keep HRR advisory until a protocol and validation dataset justify a stronger role."),
        ("Conditioning progression and recovery", "A low recovery signal may ease today’s prescription without destroying the earned conditioning level.", "`conPrescription` applies a daily adjustment while `conAdapt` banks earned level from the completed session.", "Using the same low score to both reduce today and erase the banked level can double-penalise the athlete.", "Earn a level, prescribe on a low-recovery day, complete an on-target adjusted session, and inspect the next level.", "Separate what the athlete earned from what today is wise to attempt."),
        ("Overcooked sessions", "High-zone exposure or excessive effort can make completion inadequate even when minutes are high.", "The conditioning logic has an overcooked guard based on high-zone share and effort context.", "Volume alone cannot distinguish productive work from a session that overspent capacity.", "Hold duration constant while varying high-zone share and RPE; check progress, hold, and miss paths.", "The engine should describe overcooked evidence without diagnosing the cause."),
        ("Strength-conditioning interference", "A hybrid engine must account for the interaction of hard conditioning and strength exposures without pretending to model every mechanism.", "Whole-athlete state counts recent hard sessions and lower-body hard sessions; balance and conditioning modules remain separate.", "An easy-looking lift after a hard interval may be a poor place to escalate load, while one bad lift should not erase all aerobic work.", "Use paired-week fixtures with varied ordering, density, and modality and inspect domain-specific decisions.", "Use a conservative context signal and preserve domain-specific evidence."),
        ("Minimum viable conditioning", "A short session can preserve the habit and purpose even when the full prescription is not appropriate.", "Whole-athlete state contains a time-limited constraint and the acceptance matrix calls for a minimum viable session.", "Treating a short day as failure encourages make-up work and recovery debt.", "Limit available minutes to several bands and verify that the output states what purpose is retained.", "Reduce optional work before discarding the training purpose."),
        ("Conditioning substitutions", "A substitute can preserve an energy-system purpose while remaining distinct data.", "The product architecture includes multiple modalities and the dossier requires explicit substitution records.", "Merging a bike substitute into a running trend can create a false baseline.", "Substitute modalities with identical nominal duration and compare trend keys and explanations.", "Substitution is an intentional change of evidence, not a silent equivalence."),
    ]),
    ("Chapter 5 — Readiness, HRV, and whole-athlete state", [
        ("Readiness is a state estimate", "Readiness should summarise available observations with data quality rather than pretend to measure a hidden physiological truth.", "Whole-athlete state returns score, band, signals, rationale, and data quality.", "A single composite score can hide whether the problem is sleep, soreness, stress, illness, or missing data.", "Create the same score from different signal combinations and inspect the rationale and confidence.", "Show the inputs and limitations alongside the band."),
        ("HRV as advisory context", "HRV can inform a decision but should not independently prescribe load or create a diagnosis.", "State tests assert that HRV alone does not change readiness or create a pain gate.", "A low value can be caused by measurement conditions, illness, stress, or normal variation.", "Hold direct performance constant while varying HRV and verify advisory-only behaviour.", "If HRV changes an action, the decision must also state the corroborating evidence and protocol."),
        ("Missingness as information", "A missing wearable signal is not a zero, a normal value, or a negative result.", "The state model distinguishes unknown and data quality bands; conditioning excludes absent zone data.", "Default values manufacture confidence and can drive a false reduction or progression.", "Remove one signal at a time and then all signals; compare decisions and explanations.", "Represent missingness explicitly through every boundary."),
        ("Sleep and stress context", "Sleep and life stress may alter today’s capacity but do not identify a local movement limitation.", "State derives signals and rationale from sleep, stress, soreness, life load, and training density.", "A general readiness reduction can be applied too broadly to a pain-specific movement or too narrowly to a systemic issue.", "Vary one context domain at a time and inspect whether strength and conditioning capacity remain separate.", "Context should constrain the session purpose proportionally and transparently."),
        ("Recovery debt", "Accumulated context can justify avoiding another hard session without becoming a permanent label.", "The state model estimates debt from recent observations and hard-session density.", "A debt score can create a feedback loop where reduced training prevents new evidence and the score never clears.", "Run hard weeks, recovery weeks, missing-observation weeks, and a single unusually stressful day.", "Debt needs decay, observation windows, and a clear path back to normal."),
        ("Pain hard constraint", "Pain is not an ordinary fatigue score and must outrank progression logic.", "The state package emits `pain_hold_active` as a hard constraint with an adjustment to modify or stop.", "If the flag is only displayed but not consumed, the system can still offer normal escalation.", "Raise the flag through the canonical input, request a session, and verify that the final athlete action is held, modified, or blocked.", "The consumer integration test is a release blocker."),
        ("Illness hard constraint", "Illness status requires a return-to-training process rather than a generic low-readiness adjustment.", "State emits `illness_flag_active` as a hard constraint and explicitly avoids diagnosis.", "An illness flag that only lowers load may still recommend high-intensity work.", "Test active, suspected, and clear statuses across mobile, coach, and sync paths.", "Illness must select a safe return state and preserve the distinction from ordinary tiredness."),
        ("Low energy availability context", "Nutrition facts can inform training context without becoming a training app’s food prescription.", "Nutrition context contains logged intake and estimated expenditure but no calorie target or macro instruction.", "A target passed into state can make training secretly prescribe nutrition through a side door.", "Attempt to pass target-like fields, test sparse logging, and verify only soft training constraints are emitted.", "Keep nutrition facts observational and keep food instruction in the nutrition engine."),
        ("Capacity by domain", "Strength and conditioning may be affected differently by the same context.", "The state snapshot contains separate capacity bands for overall, strength, and conditioning.", "One overall band can flatten a useful choice such as easy aerobic work while holding heavy lower-body loading.", "Create domain-specific training facts and verify the selector can choose a preserved purpose.", "Use the least restrictive safe domain-specific action."),
        ("The integration gap", "A well-typed state package is not a safety feature until a caller reads it.", "The README’s implementation note says pain and illness flags were not consumed downstream at the audit point.", "The engine may continue to offer normal prescriptions despite a hard flag.", "Build a full path test from state derivation to actual session card, not only a unit test of `deriveAthleteState`.", "Close this gap before describing the project as finished."),
    ]),
    ("Chapter 6 — Pain, illness, layoff, and calibration", [
        ("Pain hold is movement-specific", "A pain event should protect the affected movement or pattern without unnecessarily blocking unrelated safe work.", "Conditioning holds are keyed by exact format and modality, while whole-athlete state can represent a broader safety hold.", "A global stop can be overly restrictive; a local hold can be dangerously narrow when symptoms are systemic.", "Raise a lower-back pain hold, request upper-body and lower-body sessions, and verify the intended scope.", "Scope must be declared and the explanation must name the affected area or pattern."),
        ("Pain stop versus hard effort", "A set stopped for pain cannot be scored as a normal miss.", "Conditioning completion distinguishes `pain_stop`; pain holds persist until acknowledgement or a non-pain result according to the current rule.", "Mixing pain with fatigue contaminates progression streaks and can reward pushing through symptoms.", "Stop an interval for pain, stop a lift for effort, and compare their state transitions and stored evidence.", "Pain is a separate pathway from the performance controller."),
        ("Acknowledgement semantics", "A hold should persist until a meaningful acknowledgement or a valid state change, not merely until reload.", "The conditioning hold is keyed to an acknowledgement record and exact format/modality.", "A stale UI flag or an overly eager auto-clear can re-expose the athlete to the same problematic stimulus.", "Reload offline, sync across devices, acknowledge on one device, and inspect the other device’s state.", "Persist the hold and the acknowledgement event with timestamps and scope."),
        ("Red flags and clinical review", "Some inputs exceed the authority of an autonomous training engine.", "The state model says illness is not diagnosis and the dossier requires clinical review for red-flag symptoms.", "A confident exercise substitution can sound like medical advice.", "Test configured red flags and verify that the interface stops at handoff language rather than prescribing around them.", "Autonomy must include a principled abstention state."),
        ("Short training gaps", "A missed day should not be treated like detraining.", "The repository’s adaptive design distinguishes insufficient history and calibration from ordinary progression.", "A universal time-off percentage overreacts to scheduling noise.", "Use one day, three days, one week, and a longer break with otherwise identical history.", "Gap length and modality-specific context should decide whether calibration is needed."),
        ("Longer returns", "After a meaningful break, the engine should gather current evidence before chasing old loads.", "The dossier’s calibration state and adaptive contracts provide a place for conservative re-entry.", "Using the previous peak directly can create avoidable failure; using a huge arbitrary reduction can waste information.", "Return at several fractions of the old anchor, record completion/RPE/symptoms, and verify deterministic calibration exit.", "Calibration is a measurement protocol, not a universal detraining equation."),
        ("Illness return", "Return from illness is not the same as a normal layoff because symptoms and systemic tolerance matter.", "Whole-athlete state carries illness status and separates it from readiness signals.", "An athlete can be recovered from a scheduling gap but still not ready for intensity after illness.", "Keep illness active across a gap, then clear it and test the first two exposures.", "Require a clear state transition and conservative return evidence."),
        ("Equipment changes during calibration", "A new bar, machine, or plate set changes comparability even if the exercise name is unchanged.", "Plate math and exercise identity provide the foundation for equipment-aware comparisons.", "A familiar load on unfamiliar equipment may not represent the same exposure.", "Change equipment identity in the middle of a return sequence and inspect comparability and anchoring.", "Calibration should restart or widen uncertainty when equipment changes materially."),
        ("Substitution after pain", "A safe alternative must preserve purpose while remaining separate evidence.", "The product contract distinguishes substitution from the original movement trend.", "Silently assigning a substitute’s success to the painful movement can falsely clear the hold.", "Substitute, complete, and inspect both the original hold and the alternative trend.", "A substitute does not prove the original movement is ready."),
        ("Exit criteria", "States need explicit exit conditions or they become permanent or arbitrary.", "The adaptive contract includes action and safety fields but the consumer must define transition rules.", "A hidden timer can exit calibration or a pain hold without new evidence.", "Document and test every entry and exit edge, including missing data and sync conflicts.", "A state is closed only when entry, persistence, exit, and audit are all specified."),
    ]),
    ("Chapter 7 — Data quality, persistence, and sync", [
        ("Raw facts versus derived decisions", "A derived load or readiness band should never replace the observations that made it possible.", "The engine contains database sanitisation, restore, cloud mapping, and merge boundaries.", "Without raw facts, a later rule cannot replay or challenge an earlier decision.", "Delete a derived field and replay from raw session and context facts; compare the result.", "Persist observations, rule version, and derived output together."),
        ("Idempotent session completion", "Offline retries must not create duplicate exposures.", "The README highlights live sync and the acceptance matrix requires stable identifiers and idempotent replay.", "Duplicate success records can falsely satisfy a confirmation count.", "Submit the same completed session repeatedly from an offline queue and inspect history and progression.", "Idempotency is an adaptive-safety property because streaks depend on count."),
        ("Out-of-order completion", "A late-restored session must not overwrite a newer earned load.", "`liftAdapt` checks stored timestamps before replacing the progression map.", "A backup or delayed sync can rewind an athlete to an older opener.", "Complete sessions in chronological and reverse arrival order and compare the final map and audit log.", "Use event time and arrival time separately."),
        ("Sanitising numeric inputs", "Invalid numbers must fail safely at the boundary before they reach a progression formula.", "`plates.ts` and numeric helpers reject nonfinite values and clamp or sanitise where appropriate.", "Infinity, NaN, empty strings, or negative loads can create false huge jumps or broken UI.", "Fuzz load, reps, RPE, increments, units, and timestamps through the public functions.", "Sanitisation must preserve the fact that input was invalid rather than quietly inventing a value."),
        ("Merge authority", "Different fields may have different owners and conflict policies.", "The repo separates coach assignment and athlete result concepts and provides merge/restore tests.", "A single record-level last-write-wins policy is too coarse for adaptive data.", "Conflict target, actual set, note, pain flag, and derived output independently.", "Merge by declared field authority and retain conflict metadata."),
        ("Tombstones and deletions", "A deleted assignment or movement must not reappear from an old device.", "The acceptance matrix calls for deletion tombstones and offline conflict protection.", "Resurrection can expose an outdated or unsafe prescription.", "Delete on one device, replay an old queue on another, and sync in both orders.", "Deletion is a state transition with provenance, not absence of a row."),
        ("Provider data boundaries", "Wearable credentials and provider secrets must never be treated as ordinary client data.", "The README says WHOOP and Concept2 functions are server-side and tokens are encrypted in Netlify Blobs.", "A cached browser response or service-worker asset can leak credentials or stale provider data.", "Inspect browser storage, service-worker cache, function responses, and CSP behaviour in a deployed smoke test.", "Security boundaries are part of adaptive correctness."),
        ("Stale recovery data", "A yesterday’s recovery value must not silently control today’s same-day prescription.", "The code distinguishes session-captured recovery from late current data and the dossier requires stale markers.", "Retroactive readings can change a verdict after the athlete has already trained.", "Use timestamps around midnight, delayed sync, and a late WHOOP update.", "Persist the input timestamp and apply freshness rules at the decision point."),
        ("Audit record completeness", "A user or reviewer must be able to reconstruct why a number changed.", "Adaptive explanations expose reason codes, confidence, safety state, and data limitations.", "A final load without anchor, actual jump, equipment step, or evidence cannot be audited.", "Snapshot a progression event and verify replay with the same engine version.", "The audit record is a first-class product output, not debug logging."),
        ("Reset and restore", "Local reset must be explicit and restore must fail closed on corruption.", "The repository includes restore and storage modules and the README describes local-first behaviour.", "A destructive reset or partial restore can erase evidence or create a mixed-version state.", "Corrupt backups, cancel resets, restore old versions, and verify current state remains intact on failure.", "Make recovery operations reversible where practical and visibly confirmed."),
    ]),
    ("Chapter 8 — Explanations, confidence, and human control", [
        ("Reason codes are contracts", "A stable reason-code vocabulary lets UI, analytics, and audits agree on meaning.", "`adaptive/types.ts` defines a closed reason-code set and explanation shape.", "Free-form strings drift, become untestable, and can imply a stronger claim than the engine supports.", "Exhaustively switch over reason codes and ensure every action has a non-empty note.", "Keep machine reason and human note separate but generated from the same decision."),
        ("Confidence is not probability of success", "A high-confidence engine decision should not be rendered as a guarantee that the athlete will succeed.", "The explanation contract distinguishes confidence levels and data limitations rather than claiming a calibrated probability.", "A percentage badge can be interpreted as a physiological forecast without validation.", "Inspect copy for low, medium, and high confidence under both rich and sparse data.", "Describe confidence as evidence quality and rule certainty, not outcome certainty."),
        ("Safety state versus confidence", "A decision can be low-confidence but still safe to try, or high-confidence that the correct action is to block.", "The contract carries confidence and safety state as separate fields.", "Combining them into one score can make a hard stop look optional.", "Create low-confidence approved, high-confidence reduced, and blocked cases and verify UI hierarchy.", "Safety state outranks persuasive copy and user motivation."),
        ("The explanation must match the number", "A reason that describes an eased load while the field shows an earned load is a functional defect.", "`explainWorkingWeight` and `explainConPrescription` consume already-computed objects rather than recomputing them.", "Parallel calculation in a view can drift after a constant changes.", "Compare all visible numbers and notes against one fixture across mobile and web.", "Make the engine result the single source for value and explanation."),
        ("Hold explanations", "A hold should tell the athlete what was protected and what would unlock progression.", "Adaptive reason codes include insufficient history, mixed results, and already-earned load.", "‘No change’ sounds arbitrary and encourages manual escalation.", "Trigger each hold reason and verify a specific next-evidence sentence.", "The next useful action is part of the hold decision."),
        ("Reduction explanations", "A reduction must name the changed lever and avoid blaming the athlete.", "The contract differentiates reduce load, reduce volume, and deload actions.", "A generic ‘recovery low’ message cannot tell whether load, sets, or intensity was changed.", "Compare strength, conditioning, pain, illness, and time-limited reductions.", "Explain purpose preserved, lever changed, and recheck condition."),
        ("Coach override provenance", "Human judgement can be correct without being an automatic engine result.", "The dossier and repository boundary support coach workspace ownership and explicit decision sources.", "A coach override can later be mistaken for validated algorithmic evidence.", "Override a progression, sync it to mobile, and inspect source, author, timestamp, and subsequent outcome.", "Store automatic, coach-authored, athlete-adjusted, and clinical-review sources distinctly."),
        ("Athlete agency", "An adaptive offer should support the athlete’s judgement without turning a safe block into a challenge.", "The lift path keeps the field typeable while hard safety states are defined separately.", "A motivational override can undermine the exact safety layer the engine exists to protect.", "Try editing an eased offer, a hold, a pain block, and a clinical-review state.", "Permit agency inside safe bounds and make prohibited bypasses clear and respectful."),
        ("Accessibility of explanations", "A safe explanation must be understandable under time pressure and available to different users.", "The web/mobile verification surface includes UI and reachability checks; the dossier adds accessible touch and reduced-motion requirements.", "A color-only red state or a hidden tooltip can make the safety decision invisible.", "Test screen reader labels, touch targets, contrast, reduced motion, and offline rendering.", "Safety information must survive presentation changes."),
        ("Audit view versus athlete view", "The athlete needs clarity; the reviewer needs detail; both should derive from one record.", "The typed explanation includes concise note plus structured limitations and prescription.", "Showing raw JSON to athletes is unusable, while hiding it from reviewers prevents audit.", "Render the same decision in compact athlete, coach, and audit formats and compare semantic fields.", "Use layered presentation without creating layered truth."),
    ]),
    ("Chapter 9 — Testing strategy and release evidence", [
        ("Unit tests as local proofs", "Pure functions are the right place to prove arithmetic, parsing, and state transitions.", "The engine contains focused tests for lift, autoregulation, conditioning, plates, cloud, restore, parity, and adaptive decisions.", "A green local test can still miss caller reachability and deployment semantics.", "Map each rule to at least one direct unit fixture and one boundary test.", "Unit tests prove local contracts; they do not prove the product path."),
        ("Property-based numeric testing", "Load and equipment arithmetic has a large edge space that hand-picked examples do not cover.", "Numeric helpers and plate math expose deterministic public functions suitable for fuzzing.", "Rare NaN, unit, tie, or rounding cases can create large real-world jumps.", "Generate finite and invalid anchors, increments, plate sets, and unit conversions and assert invariants.", "Test monotonicity, boundedness, achievable reporting, and no false hit claims."),
        ("Scenario replay", "Adaptive behaviour is a sequence problem, not only a single-call problem.", "The engine persists progression maps and has session history and adaptive exposure selection.", "A rule can pass isolated tests while oscillating, compounding reductions, or forgetting a hold over time.", "Replay weeks containing successes, misses, illness, pain, gaps, substitutions, and device changes.", "Store scenario fixtures as decision timelines with expected states."),
        ("Integration tests for hard constraints", "The most important missing proof is that state output reaches the final prescription.", "Whole-athlete tests currently prove derivation; README notes downstream consumption is incomplete.", "A unit-perfect safety state can coexist with an unsafe logger offer.", "Run state derivation, session selection, opening-load resolution, UI rendering, and sync as one test.", "Do not close the project until this path is green."),
        ("Reachability checks", "Dead code is not product behaviour even when it is tested.", "The root scripts include reachability checks and the repo has removed unused surfaces/packages.", "A helper can remain green in isolation while no route imports it.", "Trace adaptive functions from app entry points and fail if the intended consumer is absent.", "Every safety-critical function needs a live caller test."),
        ("Parity checks", "Coach and athlete surfaces must agree on the same facts and decisions.", "The root scripts include behavioural, visual, harness, and mobile parity checks.", "A coach can see one load while the athlete’s logger opens another.", "Publish a target, log a result, sync, and compare both surfaces at each state.", "Parity is semantic first and visual second."),
        ("Offline and reload tests", "The safety decision must survive the conditions in which an athlete actually trains.", "README and scripts mention offline reload and service-worker behaviour.", "A stale cached app can offer an old rule or lose a pain hold.", "Go offline before and after state changes, reload, update the service worker, and inspect the visible decision.", "Offline behaviour must be an explicit degraded-data state, not an accident."),
        ("Deployment smoke", "A deployed function, CSP policy, and service worker can alter the effective product.", "The root verification includes build, CSP, React smoke, and deploy smoke.", "Local tests cannot see production headers, route wiring, or cache behaviour.", "Run deployment smoke against the release URL with browser section failures treated as failures.", "Archive deployed commit, environment mode, and smoke results."),
        ("Mutation testing of guards", "A test suite should fail when a critical guard is removed.", "The code contains deliberate guards for warm-ups, nonfinite plates, missing HR data, and pain/illness separation.", "A test that only checks output on happy paths may not protect the guard itself.", "Temporarily remove or invert each guard and verify a test fails.", "Use mutation-style review for safety-critical conditions."),
        ("Release evidence bundle", "A close-out release needs a compact evidence package that a future reviewer can reopen.", "The project already preserves evidence bundles, design notes, source links, and acceptance matrices.", "Without commit, test, and rule-version metadata, later review becomes archaeology.", "Package source snapshot, test output, decision schema, known gaps, and rollback instructions.", "Freeze the baseline and reopen only for a named trigger."),
    ]),
    ("Chapter 10 — Product closure, governance, and future validation", [
        ("What ‘closed’ means", "Closing the research loop means the product rule is explicit and buildable, not that every uncertainty has disappeared.", "The dossier records evidence status, provisional constants, implementation rules, and release gates.", "Endless research can become avoidance of a finite engineering decision.", "Review each open question and assign it to implemented rule, validation study, or release blocker.", "Close the question when the decision, provenance, and test are assigned."),
        ("The 2.5% default", "A conservative default can be useful without being promoted to a scientific optimum.", "The dossier chooses 2.5% of a stable opening load inside the broader ACSM practical range.", "Users may read a product default as a universally validated prescription.", "Display the rule status in developer and coach documentation and replay across equipment scales.", "Label it an engineering heuristic and tune it from observed bounded outcomes."),
        ("The 5% reactive reduction", "A repeated-decline heuristic can protect the next exposure while avoiding a compounding collapse.", "The dossier and acceptance matrix specify reduction from the last successful anchor, separate from within-session back-off.", "A flat or compounded reduction can punish one bad day or reduce too far on a light movement.", "Compare one miss, repeated misses, pain miss, gap-separated misses, and anchor changes.", "Keep the heuristic bounded, auditable, and reopenable from incident data."),
        ("Confirmation count", "Requiring more than one comparable success filters noise but delays progression.", "The product lock uses two comparable confirmations as a debounce heuristic; adaptive strength code also requires a minimum history before suggestions.", "Too little confirmation creates volatility; too much creates stagnation.", "Simulate fast responders, slow responders, beginners, and sparse logs with the same rule.", "Treat confirmation as a product tuning parameter, not a claim of universal physiology."),
        ("Outcome validation", "The engine needs product validation that is separate from its scientific rationale.", "The evidence dossier separates adherence, comprehension, safety-state integrity, performance, and superiority claims.", "More completed workouts or improved strength cannot by themselves prove causal superiority.", "Define pre/post and comparison analyses before collecting outcomes.", "Report usefulness honestly and do not overclaim experimental efficacy."),
        ("Incident review", "A safety incident should reopen the relevant rule rather than trigger a vague promise to improve AI.", "The engine has explicit states, reason codes, raw facts, and versioned decisions suitable for reconstruction.", "Without an incident taxonomy, the project may change constants without understanding the failure.", "Classify incident source, data quality, state precedence, UI bypass, sync, and human override.", "Reproduce first, patch second, retest the neighbouring states third."),
        ("Rule change governance", "Changing a constant changes athlete behaviour and must be versioned like a schema migration.", "The operational protocol requires engine version, replay fixtures, risk review, and old-decision reproducibility.", "A silent constant change makes historical decisions impossible to explain.", "Change one constant in a fixture set and inspect old and new outputs under their recorded versions.", "Release rule changes intentionally and retain the old interpretation."),
        ("Human handoff", "A good engine knows when the next action belongs to a coach or clinician.", "The state and explanation contracts include blocked, held, safety, and review concepts.", "Forcing an autonomous answer can be more harmful than admitting uncertainty.", "Trigger red flags, persistent pain, complex conflict, and severe missingness and inspect the handoff.", "Abstention is a completed decision when the engine lacks authority."),
        ("Agent-assisted review discipline", "Independent review roles can increase coverage when their outputs are integrated and verified.", "The repository contains locator/reviewer agent instructions and a dispatching skill.", "Agent output can be mistaken for evidence or can duplicate work without a final integration pass.", "Require exact file/symbol findings, classify each finding, then run the full test and human review.", "Use agents as bounded reviewers; never let them become an unverified source of truth."),
        ("Final release checklist", "Closure requires a finite checklist that can be signed off.", "The repository’s `verify` script and dossier acceptance matrix provide the base.", "A project can feel finished while one consumer path or safety state remains unproven.", "Run the checklist at the release commit and record every pass, skip, failure, and owner.", "Release only when hard-constraint integration, athlete reachability, parity, and deployment smoke are closed."),
    ]),
]


def render_monograph_topic(title, focus, current, risk, test, decision, chapter_no, topic_no):
    return dedent(f'''
    ### {chapter_no}.{topic_no} {title}

    **Question.** {focus} The implementation question is not whether this sentence sounds sensible. It is whether the product can represent the relevant fact, preserve it across the session lifecycle, choose a bounded action, and explain the action without overstating the evidence. In this project, the answer must be traceable from an input field or sensor observation to an engine function, then to a persisted decision and the screen that the athlete or coach actually uses.

    **Repository reading.** {current} That evidence shows an intentional boundary, not automatically a complete feature. A named type, helper, or unit test establishes a local contract. It does not establish that the function is imported by the live caller, that its result survives offline storage, that the coach and athlete surfaces agree, or that a hard state cannot be bypassed by an alternate route. The audit therefore treats implementation, reachability, integration, and release proof as separate claims.

    **Evidence status.** The scientific conclusion for this mechanism is deliberately narrower than the product ambition. Research can support the direction of a controller, the usefulness of individualisation, the value of observing performance and symptoms, or the danger of false certainty. It usually does not validate the exact constant, confirmation count, UI wording, or threshold selected here. The product may still choose a heuristic when it is bounded, reversible, transparent, and low-risk, but the documentation must call it a heuristic and identify what would change it.

    **Failure mode.** {risk} The important distinction is between an arithmetic error, a data-quality error, a state-precedence error, and a communication error. Each can produce a different visible result even when the same underlying observation is present. A safe implementation names the failure class, retains the raw input, and avoids converting a missing or incomparable observation into apparent evidence. If the rule cannot tell whether the event was pain, fatigue, equipment failure, or incomplete logging, it should hold or ask for input rather than invent a cause.

    **State transition.** The mechanism should enter a declared state before it changes a prescription. A normal approved path may allow the planned exposure. A held path preserves the current dose while gathering evidence. A reduced path changes one lever and states what purpose remains. A calibration path gathers a new anchor after a meaningful change. A blocked or review path prevents autonomous escalation. The order matters: hard pain or illness constraints outrank ordinary performance progression; missing optional data lowers confidence; one noisy miss does not automatically become a deload.

    **Numerical discipline.** Any number shown to the athlete should retain its units, reference anchor, rounding rule, and actual physical meaning. The engine should distinguish intended percentage from achieved percentage, prescribed load from logged load, and eased offer from earned baseline. Rounding must happen at the equipment boundary, not earlier in a way that hides the real jump. When a number cannot be represented safely, the correct output is a fallback, hold, or request for setup information—not a false statement that the target was achieved.

    **User experience.** The athlete should see the next useful action, not a lecture. A progression note should identify the evidence and the single changed lever. A hold should say what was protected and what evidence would unlock a change. A reduction should say whether load, volume, density, or complexity moved. A blocked state should be respectful, unambiguous, and impossible to mistake for a motivational challenge. A coach should be able to open the underlying observations and see whether the decision was automatic, athlete-adjusted, coach-authored, or review-bound.

    **Required validation.** {test} The test must exist at the pure function level and at the boundary where the decision reaches the real product surface. Sequence tests are essential because adaptive errors often arise from a prior session, a delayed sync, a stale wearable value, or a previously active hold. The fixture should include valid data, missing data, conflicting data, a user override, and a hard safety flag. The expected result should include action, source, reason codes, safety state, and the next recheck condition.

    **Audit record.** Every decision in this area should retain the engine version, time, subject and exercise identity, session purpose, source observations, comparability status, anchor, equipment step, actual rounded change, data limitations, safety constraints, reason codes, and chosen action. The interface can summarise these fields, but a reviewer must be able to reconstruct the decision without guessing which screen or constant supplied the number. This is the difference between a transparent controller and a black-box recommendation that merely happens to be deterministic.

    **Closure judgement.** {decision} This is strong enough to freeze as the implementation position for the current release, subject to the explicit integration and verification gates in the repository audit. It should not be described as a universal law of training. It is a bounded product policy whose safety depends on truthful missingness, correct state precedence, equipment-aware arithmetic, and an honest handoff when the evidence or authority runs out.

    ''')


def build_expanded_monograph():
    out = [REPO_AUDIT]
    out.append("## Part VIII. Full mechanism-by-mechanism close-out monograph\n\n")
    out.append("This part expands the implementation lock into a searchable review record. Each subsection is intentionally framed as a decision, evidence boundary, failure analysis, state transition, validation requirement, and closure judgement. The repeated structure is a control against selective attention: every mechanism must answer the same questions before it is allowed to change training.\n\n")
    for chapter_no, (chapter_title, topics) in enumerate(MONOGRAPH_PACKS, start=1):
        out.append(f"## {chapter_title}\n\n")
        for topic_no, row in enumerate(topics, start=1):
            out.append(render_monograph_topic(*row, chapter_no, topic_no))
    out.append(dedent(r'''

    ## Part IX. Final implementation lock

    ### The deterministic controller that is now approved

    The approved baseline is a transparent controller with separate pathways. The strength pathway reads comparable working exposures, excludes warm-ups from the working anchor, evaluates completion and effort, respects equipment reality, and chooses one primary lever. The conditioning pathway keys progression to format and modality, separates earned level from today’s recovery adjustment, excludes simulated or data-empty sessions, and keeps provisional thresholds visibly provisional. The whole-athlete pathway collects context and hard constraints but does not allow a composite readiness score to impersonate a diagnosis. The pain and illness pathway outranks ordinary fatigue logic. The data boundary preserves raw facts and decision provenance. The explanation boundary returns the same number and the same reason to every surface.

    The chosen values are defaults, not laws: a 2.5% target relative to a last stable opening load; equipment-aware rounding; a 5% reactive reduction from the last successful anchor after repeated comparable deterioration; two comparable successes as a progression debounce; one miss as hold/context; repeated comparable misses as reduction or recheck; and a minimum history gate before high-confidence adaptive suggestions. These values are safe only because they are bounded and because the controller can hold, fall back, calibrate, or abstain. If the application removes those escape states, the constants become materially riskier.

    ### The finite blockers before public closure

    The project should not reopen the entire research question. It should close the finite implementation blockers. First, prove that whole-athlete pain and illness hard constraints are consumed by the actual session/prescription path and cannot be bypassed by the mobile logger or a stale synced assignment. Second, prove that the new adaptive strength output reaches the athlete’s live opening-load and explanation surface. Third, prove coach/athlete field ownership and idempotent sync with conflict fixtures. Fourth, run the complete repository verification and deployment smoke at the release commit, including browser sections that must not silently skip. Fifth, archive the decision record, test output, repository commit, known limitations, and rollback path.

    ### What would reopen the decision

    The rule should reopen only for a named trigger: a safety incident, a reproducible systematic bias, a new direct comparative trial that changes the evidence boundary, a strong product validation signal showing repeated under- or over-progression, a new modality whose measurement semantics do not fit the current controller, or a schema/security change that breaks replay. A single unusually good workout, a single low HRV value, a request for a more aggressive number, or the existence of a new general review is not by itself enough. Reopening should preserve the current baseline while the new evidence is tested.

    ### Final verdict

    THE Hybrid Engine is ready to close as a deterministic adaptive-training project once the finite integration gates are green. The science supports individualised, bounded, feedback-informed progression and cautions against false precision. The repository contains the beginnings—and in several areas the substance—of that design: authoritative lift calculations, equipment-aware feasibility, format-specific conditioning, explanation contracts, whole-athlete state, explicit missingness, and layered verification. The remaining work is not to discover a perfect percentage. It is to connect the safety state to the real consumer, prove the live path, preserve the audit trail, and release the bounded controller with its uncertainty honestly labelled.

    '''))
    return "".join(out)


EXPANDED_MONOGRAPH = build_expanded_monograph()


def main():
    report = add_generated_material(AUTHORED)
    OUT_MD.write_text(report, encoding="utf-8")
    build_docx(report)
    print(f"REPORT_WORDS={len(report.split())}")
    print(f"MARKDOWN={OUT_MD}")
    print(f"DOCX={OUT_DOCX}")

if __name__ == "__main__":
    main()
