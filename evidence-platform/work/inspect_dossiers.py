from pathlib import Path
import hashlib, json, os, shutil

from docx import Document
from pypdf import PdfReader
import pypdfium2 as pdfium
from PIL import Image, ImageDraw

BASE=Path(__file__).resolve().parents[1]
ROOT=BASE/'sources'/'original-archive'/'THE-Hybrid-System-Master-Evidence-Archive-2026-08-28'
LONG=Path('\\\\?\\'+str(ROOT)) if os.name=='nt' else ROOT
qa=BASE/'work'/'qa-dossiers'; qa.mkdir(parents=True,exist_ok=True)
found=[]
for dp,_,names in os.walk(str(LONG)):
    for name in names:
        if Path(name).suffix.lower() in {'.docx','.pdf'}:
            found.append(Path(dp)/name)
for src in found:
    shutil.copyfile(src, qa/('source-dossier'+src.suffix.lower()))

docx_path=qa/'source-dossier.docx'
pdf_path=qa/'source-dossier.pdf'
doc=Document(docx_path)
doc_text='\n'.join(p.text for p in doc.paragraphs)
(qa/'docx-text.txt').write_text(doc_text,encoding='utf-8')
reader=PdfReader(pdf_path)
pdf_text='\n'.join((p.extract_text() or '') for p in reader.pages)
(qa/'pdf-text.txt').write_text(pdf_text,encoding='utf-8')

pdf=pdfium.PdfDocument(pdf_path)
page_images=[]
for i in range(len(pdf)):
    img=pdf[i].render(scale=1.1).to_pil().convert('RGB')
    out=qa/f'pdf-page-{i+1:03}.png'; img.save(out)
    thumb=img.copy(); thumb.thumbnail((380,500)); page_images.append((i+1,thumb))
for batch_start in range(0,len(page_images),12):
    batch=page_images[batch_start:batch_start+12]
    sheet=Image.new('RGB',(4*400,3*540),'white'); draw=ImageDraw.Draw(sheet)
    for j,(num,img) in enumerate(batch):
        x=(j%4)*400+10; y=(j//4)*540+25
        draw.text((x,y-20),f'Page {num}',fill='black'); sheet.paste(img,(x,y))
    sheet.save(qa/f'contact-{batch_start//12+1:02}.png')

report={
  'docx_sha256':hashlib.sha256(docx_path.read_bytes()).hexdigest().upper(),
  'pdf_sha256':hashlib.sha256(pdf_path.read_bytes()).hexdigest().upper(),
  'docx_paragraphs':len(doc.paragraphs),'docx_tables':len(doc.tables),
  'docx_inline_shapes':len(doc.inline_shapes),'docx_text_chars':len(doc_text),
  'pdf_pages':len(reader.pages),'pdf_text_chars':len(pdf_text),
  'pdf_rendered_pages':len(page_images),
  'text_prefix_match':doc_text[:500].strip()==pdf_text[:500].strip(),
  'note':'Visual QA requires contact-sheet inspection; DOCX could not be rendered because LibreOffice is unavailable.'
}
(qa/'inspection.json').write_text(json.dumps(report,indent=2)+'\n',encoding='utf-8')
print(json.dumps(report,indent=2))
